# -*- coding: utf-8 -*-
"""
write_docx参数组合测试 - 小健 2026-06-24

测试类型:
1. 基础功能测试 - 参数组合
2. Markdown功能测试 - 各语法元素
3. 真实场景测试 - 业务文档
4. 边界测试 - 特殊情况
5. 负面测试 - 错误处理
"""

import pytest
from pathlib import Path
from docx import Document
from app.tools.document.write_docx import write_docx
from app.tools.tool_response import is_success, is_error


class TestWriteDocxBasicParams:
    """基础参数组合测试"""

    def test_empty_document(self, temp_output_dir, docx_test_data):
        """Case 1: 最小文档 - 进化: content不能为空字符串,需附带最小合法content"""
        file_path = temp_output_dir / "empty.docx"

        result = write_docx(str(file_path), content="默认文档内容")

        assert is_success(result)
        assert file_path.exists()
        assert str(file_path) in result["llm_data"]["summary"]

    def test_title_only(self, temp_output_dir, docx_test_data):
        """Case 2: 仅标题 - 进化: content不能为空字符串,需附带最小合法content"""
        test_data = docx_test_data["title_only"]
        file_path = temp_output_dir / "title_only.docx"

        result = write_docx(str(file_path), title=test_data["title"], content="标题下的正文内容")

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 1, "应该至少有1个段落"
        assert doc.paragraphs[0].text == "测试报告", "标题内容应该是'测试报告'"
        assert doc.paragraphs[0].style.name == 'Title', "标题样式应该是Title"

    def test_content_only(self, temp_output_dir, docx_test_data):
        """Case 3: 仅内容"""
        test_data = docx_test_data["content_only"]
        file_path = temp_output_dir / "content_only.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 1, "应该至少有1个段落"
        assert "正文内容" in doc.paragraphs[0].text, "内容应该包含'正文内容'"

    def test_title_and_content(self, temp_output_dir, docx_test_data):
        """Case 4: 标题+内容"""
        test_data = docx_test_data["simple"]
        file_path = temp_output_dir / "simple.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 2, "应该至少有2个段落(标题+内容)"
        assert doc.paragraphs[0].style.name == 'Title', "第一个应该是Title样式"


class TestWriteDocxMarkdownHeadings:
    """Markdown标题功能测试"""

    @pytest.mark.parametrize("level,prefix,content_key", [
        (1, "# ", "h1"),
        (2, "## ", "h2"),
        (3, "### ", "h3"),
        (4, "#### ", "h4"),
        (5, "##### ", "h5"),
    ])
    def test_single_heading_level(self, temp_output_dir, level, prefix, content_key):
        """测试单个标题级别"""
        file_path = temp_output_dir / f"heading_{content_key}.docx"
        content = f"{prefix}{content_key}标题内容"

        result = write_docx(str(file_path), title=None, content=content)

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 1, "应该至少有1个段落"
        expected_style = f'Heading {level}'
        assert doc.paragraphs[0].style.name == expected_style, f"应该是{expected_style}样式,实际是{doc.paragraphs[0].style.name}"

    def test_all_heading_levels(self, temp_output_dir, docx_test_data):
        """测试所有级别标题"""
        test_data = docx_test_data["headings"]
        file_path = temp_output_dir / "all_headings.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        heading_paras = [p for p in doc.paragraphs if 'Heading' in p.style.name]
        assert len(heading_paras) >= 5, f"应该至少有5个标题段落,实际有{len(heading_paras)}"

    def test_consecutive_headings(self, temp_output_dir):
        """测试连续标题(无内容)"""
        file_path = temp_output_dir / "consecutive_headings.docx"
        content = """# 标题1
## 标题2
### 标题3
#### 标题4
##### 标题5
"""
        result = write_docx(str(file_path), content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        assert len(doc.paragraphs) == 5, f"应该有5个段落,实际有{len(doc.paragraphs)}"
        assert doc.paragraphs[0].style.name == 'Heading 1'
        assert doc.paragraphs[1].style.name == 'Heading 2'
        assert doc.paragraphs[2].style.name == 'Heading 3'
        assert doc.paragraphs[3].style.name == 'Heading 4'
        assert doc.paragraphs[4].style.name == 'Heading 5'


class TestWriteDocxMarkdownLists:
    """Markdown列表功能测试"""

    def test_unordered_list_dash(self, temp_output_dir, docx_test_data):
        """无序列表 - 短横线(-)"""
        test_data = docx_test_data["unordered_list_dash"]
        file_path = temp_output_dir / "unordered_dash.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        list_paras = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(list_paras) == 4, f"应该有4个无序列表项,实际有{len(list_paras)}"

    def test_unordered_list_asterisk(self, temp_output_dir, docx_test_data):
        """无序列表 - 星号(*)"""
        test_data = docx_test_data["unordered_list_asterisk"]
        file_path = temp_output_dir / "unordered_asterisk.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        list_paras = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(list_paras) == 3, f"应该有3个无序列表项,实际有{len(list_paras)}"

    def test_ordered_list_sequential(self, temp_output_dir):
        """有序列表 - 连续数字(1. 2. 3.)"""
        file_path = temp_output_dir / "ordered_sequential.docx"
        content = """1. 第一项
2. 第二项
3. 第三项
"""
        result = write_docx(str(file_path), content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        list_paras = [p for p in doc.paragraphs if p.style.name == 'List Number']
        assert len(list_paras) == 3, f"应该有3个有序列表项,实际有{len(list_paras)}"

    def test_ordered_list_arbitrary_numbers(self, temp_output_dir, docx_test_data):
        """有序列表 - 任意数字前缀(10. 99.)"""
        test_data = docx_test_data["ordered_list"]
        file_path = temp_output_dir / "ordered_arbitrary.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)

        doc = Document(str(file_path))
        list_paras = [p for p in doc.paragraphs if p.style.name == 'List Number']
        assert len(list_paras) == 5, f"应该有5个有序列表项,实际有{len(list_paras)}"

    def test_mixed_lists(self, temp_output_dir, docx_test_data):
        """混合列表 - 无序+有序"""
        test_data = docx_test_data["mixed_lists"]
        file_path = temp_output_dir / "mixed_lists.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)

        doc = Document(str(file_path))
        bullet_paras = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        number_paras = [p for p in doc.paragraphs if p.style.name == 'List Number']
        assert len(bullet_paras) >= 3, f"应该至少有3个无序列表项,实际有{len(bullet_paras)}"
        assert len(number_paras) >= 3, f"应该至少有3个有序列表项,实际有{len(number_paras)}"


class TestWriteDocxTables:
    """Markdown表格和table_data功能测试"""

    def test_markdown_table_basic(self, temp_output_dir):
        """Markdown表格 - 基础格式"""
        file_path = temp_output_dir / "table_basic.docx"
        content = """# 数据表格

| 项目 | 数值 | 占比 |
|------|------|------|
| A | 100 | 40% |
| B | 150 | 60% |

## 结论

表格已生成."""

        result = write_docx(str(file_path), content=content)

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.tables) == 1, f"应该有1个表格,实际有{len(doc.tables)}"

        table = doc.tables[0]
        assert len(table.rows) == 3, f"表格应该有3行,实际有{len(table.rows)}行"
        assert len(table.columns) == 3, f"表格应该有3列,实际有{len(table.columns)}列"

        assert table.rows[0].cells[0].text == "项目"
        assert table.rows[0].cells[1].text == "数值"
        assert table.rows[0].cells[2].text == "占比"

    def test_markdown_table_with_content(self, temp_output_dir):
        """Markdown表格 - 混合内容(标题+段落+表格+列表)"""
        file_path = temp_output_dir / "table_mixed.docx"
        content = """# 报告标题

这是介绍段落.

## 数据统计

| 名称 | 数量 |
|------|------|
| 项目A | 100 |
| 项目B | 200 |

## 分析要点

- 要点1:数据A占33%
- 要点2:数据B占67%

## 详细步骤

1. 第一步
2. 第二步
3. 第三步"""

        result = write_docx(str(file_path), content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        assert len(doc.tables) == 1, "应该有1个表格"

        heading_paras = [p for p in doc.paragraphs if 'Heading' in p.style.name]
        assert len(heading_paras) >= 4, f"应该至少有4个标题,实际有{len(heading_paras)}"

        bullet_paras = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(bullet_paras) == 2, f"应该有2个无序列表项,实际有{len(bullet_paras)}"

        number_paras = [p for p in doc.paragraphs if p.style.name == 'List Number']
        assert len(number_paras) == 3, f"应该有3个有序列表项,实际有{len(number_paras)}"

    def test_table_data_parameter(self, temp_output_dir):
        """table_data参数 - 纯表格文档"""
        file_path = temp_output_dir / "table_data.docx"
        table_data = [
            ["姓名", "年龄", "城市"],
            ["张三", "25", "北京"],
            ["李四", "30", "上海"],
            ["王五", "28", "广州"]
        ]

        result = write_docx(str(file_path), title="人员信息表", table_data=table_data)

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.tables) == 1, f"应该有1个表格,实际有{len(doc.tables)}"

        table = doc.tables[0]
        assert len(table.rows) == 4, f"表格应该有4行,实际有{len(table.rows)}行"
        assert len(table.columns) == 3, f"表格应该有3列,实际有{len(table.columns)}列"

        assert table.rows[0].cells[0].text == "姓名"
        assert table.rows[1].cells[0].text == "张三"
        assert table.rows[2].cells[1].text == "30"

    def test_table_data_with_title(self, temp_output_dir):
        """table_data参数 - 带标题的表格文档"""
        file_path = temp_output_dir / "table_with_title.docx"
        table_data = [
            ["产品", "销量", "金额"],
            ["产品A", "1000", "10000"],
            ["产品B", "800", "8000"]
        ]

        result = write_docx(str(file_path), title="销售统计", table_data=table_data)

        assert is_success(result)

        doc = Document(str(file_path))
        assert doc.paragraphs[0].style.name == 'Title', "第一个应该是Title样式"
        assert len(doc.tables) == 1, "应该有1个表格"

    def test_content_table_data_both_render(self, temp_output_dir):
        """content与table_data并存时均渲染(非互斥)"""
        file_path = temp_output_dir / "mutex.docx"

        result = write_docx(
            str(file_path),
            title="测试",
            content="# 有content\n\n段落内容",
            table_data=[["A", "B"], ["C", "D"]]
        )

        assert is_success(result)

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 2, "应该有段落(来自content)"
        assert len(doc.tables) == 1, "table_data应渲染为1个表格"
        table = doc.tables[0]
        assert table.rows[0].cells[0].text == "A"

    def test_markdown_table_multiple(self, temp_output_dir):
        """Markdown表格 - 多个表格"""
        file_path = temp_output_dir / "tables_multiple.docx"
        content = """# 多表格文档

## 表格1

| A | B |
|---|---|
| 1 | 2 |

中间段落.

## 表格2

| C | D |
|---|---|
| 3 | 4 |"""

        result = write_docx(str(file_path), content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        assert len(doc.tables) == 2, f"应该有2个表格,实际有{len(doc.tables)}"


class TestWriteDocxRealScenarios:
    """真实场景测试"""

    def test_tech_report(self, temp_output_dir, docx_test_data):
        """场景1: 技术报告"""
        test_data = docx_test_data["tech_report"]
        file_path = temp_output_dir / "tech_report.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()
        assert result["llm_data"]["status"]["exec_code"] == "success"

        doc = Document(str(file_path))
        heading_paras = [p for p in doc.paragraphs if 'Heading' in p.style.name]
        assert len(heading_paras) >= 3, f"技术报告应该至少有3个标题,实际有{len(heading_paras)}"

    def test_meeting_minutes(self, temp_output_dir, docx_test_data):
        """场景2: 会议纪要"""
        test_data = docx_test_data["meeting_minutes"]
        file_path = temp_output_dir / "meeting_minutes.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 5, f"会议纪要应该至少有5个段落,实际有{len(doc.paragraphs)}"

    def test_from_register_examples(self, temp_output_dir):
        """场景3: register.py中的examples"""
        file_path1 = temp_output_dir / "report.docx"
        result1 = write_docx(str(file_path1), title="测试报告", content="这是测试内容")
        assert is_success(result1)

        doc1 = Document(str(file_path1))
        assert len(doc1.paragraphs) >= 2, "应该至少有2个段落"

        file_path2 = temp_output_dir / "report_structured.docx"
        content = "# 第一章\n\n正文内容\n\n## 第二节\n\n- 列表项\n- 列表项"
        result2 = write_docx(str(file_path2), title="结构化报告", content=content)
        assert is_success(result2)

        doc2 = Document(str(file_path2))
        heading_paras = [p for p in doc2.paragraphs if 'Heading' in p.style.name]
        assert len(heading_paras) >= 2, f"应该至少有2个标题,实际有{len(heading_paras)}"


class TestWriteDocxBoundary:
    """边界测试"""

    def test_special_chars(self, temp_output_dir, docx_test_data):
        """边界1: 特殊字符"""
        test_data = docx_test_data["special_chars"]
        file_path = temp_output_dir / "special_chars.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        has_special = any('<' in p.text or '>' in p.text for p in doc.paragraphs)
        assert has_special, "应该包含特殊字符<或>"

    def test_long_content(self, temp_output_dir, docx_test_data):
        """边界2: 长内容(100行)"""
        test_data = docx_test_data["long_content"]
        file_path = temp_output_dir / "long_content.docx"

        result = write_docx(str(file_path), title=test_data["title"], content=test_data["content"])

        assert is_success(result)
        assert file_path.exists()

        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 100, f"应该至少有100个段落,实际有{len(doc.paragraphs)}"

    def test_chinese_only(self, temp_output_dir):
        """边界3: 纯中文"""
        file_path = temp_output_dir / "chinese_only.docx"
        content = "这是纯中文内容测试.包含中文标点:,.,!?"

        result = write_docx(str(file_path), title="中文标题", content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        has_chinese = any('中文' in p.text for p in doc.paragraphs)
        assert has_chinese, "应该包含中文内容"

    def test_empty_lines(self, temp_output_dir):
        """边界4: 多空行"""
        file_path = temp_output_dir / "empty_lines.docx"
        content = """第一段


第二段



第三段
"""
        result = write_docx(str(file_path), content=content)

        assert is_success(result)

        doc = Document(str(file_path))
        assert len(doc.paragraphs) == 3, f"应该有3个段落(空行被跳过),实际有{len(doc.paragraphs)}"

    def test_whitespace_only(self, temp_output_dir):
        """边界5: 仅空白字符 - 空白content(len>0)允许,渲染为空文档"""
        file_path = temp_output_dir / "whitespace.docx"
        content = "   \n\n   \n   "

        result = write_docx(str(file_path), content=content)

        assert is_success(result)
        assert file_path.exists()

    def test_overwrite_existing(self, temp_output_dir):
        """边界6: 覆盖已存在文件"""
        file_path = temp_output_dir / "overwrite.docx"

        result1 = write_docx(str(file_path), title="第一次", content="内容1")
        assert is_success(result1)

        doc1 = Document(str(file_path))
        first_content = doc1.paragraphs[0].text

        result2 = write_docx(str(file_path), title="第二次", content="内容2")
        assert is_success(result2)

        doc2 = Document(str(file_path))
        assert doc2.paragraphs[0].text != first_content, "文件内容应该被覆盖"


class TestWriteDocxNegative:
    """负面测试 - 错误处理"""

    def test_invalid_path_nonexistent_drive(self):
        """负面1: 无效驱动器"""
        result = write_docx("Z:/nonexistent/path/file.docx", title="测试")

        assert is_error(result)
        assert result["llm_data"]["status"]["exec_code"] == "error"

    def test_relative_path_creates_file(self, temp_output_dir):
        """负面2: 相对路径(应能正常创建)"""
        import os
        original_cwd = os.getcwd()
        try:
            os.chdir(str(temp_output_dir))
            result = write_docx("relative.docx", title="相对路径测试")
            assert is_success(result)
            assert Path("relative.docx").exists()
        finally:
            os.chdir(original_cwd)


class TestWriteDocxLlmData:
    """llm_data结构验证"""

    def test_llm_data_success_structure(self, temp_output_dir):
        """验证成功时llm_data结构"""
        file_path = temp_output_dir / "llm_test.docx"
        result = write_docx(str(file_path), title="测试", content="内容")

        llm_data = result["llm_data"]

        assert "summary" in llm_data
        assert "action" in llm_data
        assert "status" in llm_data
        assert "duration_ms" in llm_data
        assert "metrics" in llm_data

        assert llm_data["status"]["exec_code"] == "success"
        assert llm_data["action"]["tool"] == "write_docx"
        assert llm_data["action"]["tool_zh"] == "写入Word"

    def test_llm_data_error_structure(self):
        """验证错误时llm_data结构"""
        result = write_docx("Z:/invalid/path.docx")

        llm_data = result["llm_data"]

        assert llm_data["status"]["exec_code"] == "error"
        assert "detail" in llm_data["status"]
        assert llm_data["metrics"] == {}
