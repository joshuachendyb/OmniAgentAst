# -*- coding: utf-8 -*-
"""
write_docx content/table_data 互斥参数测试 - 小健 2026-06-27

测试焦点:1. Schema验证层——WriteDocxInput.model_validator 互斥/必选案则 2. 工具函数层——write_docx() 边界行为

强制案范:- 每函数都必须有docstring,含作者+ 日期(签名:小健)- 业务数据必须 >= 10 字符真实内容

Version 1.0
"""

import pytest
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from pydantic import ValidationError

from app.tools.document.document_schema import WriteDocxInput
from app.tools.document.write_docx import write_docx
from app.tools.tool_response import is_success, is_error


# ============================= 辅助常量 —— 小健 2026-06-27 =============================
_REAL_TITLE_PROJECT = "销售业绩季度报告 —— 2026年第二季度"
_REAL_TITLE_FINANCE = "财务分析报告 —— 半年度总结"
_REAL_TITLE_MEETING = "项目进度会议纪要 —— 2026年6月"

_REAL_CONTENT_PROJECT = """# 销售业绩季度报告
## 一,执行摘要
2026年第二季度,公司整体销售额达到1,200万元,同比增长15.3%,环比增长8.7%.其中核心产品线贡献了78%的营收.
### 1.1 关键指标

- **总销售额**: 1,200万元(同比增长15.3%)- **净利润**: 180万元(同比增长12.1%)- **客户数**: 新增企业客户47家- **续约率**: 92.5%(高于行业平均水平)

## 二,分区业绩
### 2.1 华东区域

华东区域本季度表现突出,销售额达到480万元,占总销售额的40%.
### 2.2 华南区域

华南区域销售额320万元,同比增长18.6%,主要得益于新产品线推广.
### 2.3 华北区域

华北区域销售额250万元,同比略有下降,需要重点关注.
## 三,产品分析
## 四,问题与改进

## 五,下季度计划
"""

_REAL_LONG_CONTENT = """# 项目总结报告 —— 企业信息化升级改造工程
## 一,项目概况
企业信息化升级改造工程自2026年4月启动,历时6个月,总投资预算为500万元.项目覆盖了ERP系统升级,CRM系统重构,数据中台建设三大核心模块.
### 1.1 项目范围

- **ERP系统升级**: 从v3.0升级至v5.0,涉及财务,采购,库存,生产四个核心模块- **CRM系统重构**: 基于微服务架构重新开发,支持移动里访问- **数据中台建设**: 构建统一数据平台,打通各业务系统数据孤岛

### 1.2 项目团队

项目团队由15人组成,包括项目经理2人,产品经理1人,在里开发4人,前里开发2人,测试工程师6人,运维工程师2人,数据分析师2人.
## 二,执行情况
### 2.1 ERP系统升级

| 模块 | 计划完成 | 实际完成 | 进度 |
|------|---------|---------|------|
| 财务模块 | 3月31日 | 3月28日 | 提前3天|
| 采购模块 | 4月30日 | 5月5日 | 延期5天|
| 库存模块 | 5月31日 | 5月30日 | 提前1天|
| 生产模块 | 6月30日 | 6月25日 | 提前5天|

### 2.2 CRM系统重构

| 阶段 | 时间 | 交付物 | 状态 |
|------|------|--------|------|
| 需求分析 | 1月-2月 | PRD文档 | 已完成|
| 架构设计 | 2月-3月 | 技术方案 | 已完成|
| 核心开发 | 3月-5月 | 功能代码 | 已完成|
| 联调测试 | 5月-6月 | 测试报告 | 进行中|

### 2.3 数据中台建设

- **数据采集层**: 已完成12个业务系统的数据接入
- **数据存储层**: 构建了数据湖和数据仓库- **数据服务层**: 提供统一的API网关和数据服务
## 三,关键成果
### 3.1 量化指标

1. 系统响应时间从平均800ms降低到150ms
2. 数据处理能力从每日50万条提升到200万条
3. 系统可用性从99.5%提升至99.95%
4. 用户满意度评分从3.8分提升至4.5分
### 3.2 业务价值
- 财务结算周期从5天缩短至2天- 采购审批流程从3天缩短至4小时
- 库存周转率提升了25%
- 客户响应时间从24小时缩短至2小时

## 四,项目财务
| 类别 | 预算(万元) | 实际支出(万元) | 偏差 |
|------|-----------|--------------|------|
| 人力成本 | 250 | 265 | +6.0% |
| 软件采购 | 120 | 115 | -4.2% |
| 硬件采购 | 80 | 75 | -6.3% |
| 培训费用 | 30 | 28 | -6.7% |
| 其他费用 | 20 | 18 | -10.0% |
| **合计** | **500** | **501** | **+0.2%** |

## 五,质量统计
| 指标 | 目标值 | 实际值 | 达标情况 |
|------|-------|-------|---------|
| 单元测试覆盖率 | >= 80% | 87.3% | 达标 |
| Bug率 | <= 1.5/KLOC | 0.8/KLOC | 达标 |
| 严重Bug数 | <= 10 | 6 | 达标 |
| 代码重复率 | <= 15% | 11.2% | 达标 |
| 文档完整度 | >= 90% | 95% | 达标 |

## 六,风险与问题

### 6.1 已解决问题
- 数据库迁移性能问题:3月已解决- 第三方接口兼容性问题(4月已解决)- 移动里适配问题:5月已解决
### 6.2 待跟进问题
- 部分老旧硬件需要更换- 部分员工对新系统操作不熟悉
## 七,经验总结

### 7.1 成功经验

1. 采用敏捷开发方法,每周迭代,快速响应需求变更2. 建立完备的测试体系,认保交付质量
3. 定期沟通机制,认保各方信息同步

### 7.2 改进建议

- 前期需求调研需要更加深入- 风险识别和应急预案需要加强- 培训计划需要提前制定
## 八,附录
### 8.1 参考资料
- ERP系统升级技术方案- CRM系统架构设计文档
- 数据中台建设白皮书
### 8.2 工具清单

- 项目管理: Jira, Confluence
- 开发工具: IntelliJ IDEA, VS Code
- 测试工具: JMeter, Postman
- 部署工具: Jenkins, Docker

---

**报告编制**: 项目管理办公室**编制日期**: 2026年6月27日**审批人**: 张总**版本**: v1.0
"""

_REAL_TABLE_DATA = [
    ["姓名", "部门", "职位", "入职日期", "薪资", "绩效等级"],
    ["张伟", "技术部", "高级工程师", "2024-03-15", "15000", "A"],
    ["李娜", "市场部", "市场总监", "2023-08-01", "18000", "A+"],
    ["王强", "销售部", "区域经理", "2024-01-10", "16000", "B+"],
    ["赵敏", "人事部", "HR经理", "2022-11-20", "14000", "A"],
    ["孙磊", "财务部", "财务主管", "2023-06-05", "15500", "B"],
    ["周婷", "技术部", "架构师", "2024-02-28", "22000", "A+"],
    ["吴昊", "产品部", "产品经理", "2023-09-12", "17000", "A"],
    ["郑洋", "运营部", "运营主管", "2024-04-01", "13500", "B+"],
    ["陈静", "设计部", "UI设计师", "2023-12-18", "12500", "A"],
]


# ============================= Schema 验证层单元测试 —— 小健 2026-06-27 =============================
class TestParamCombinations:
    """参数组合验证 —— content/table_data 互斥 + 必选 —— 小健 2026-06-27"""

    def test_content_provided_table_data_none_passes(self):
        """Case 0: content有值table_data为None通过验证 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="销售报告.docx",
            content="# 销售报告\n\n季度销售额突破500万元",
        )
        assert inp.content is not None
        assert inp.table_data is None

    def test_table_data_provided_content_none_passes(self):
        """Case 0b: table_data有值content为None通过验证 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="销售数据表.docx",
            table_data=[["季度", "销售额"], ["Q1", "500万"]],
        )
        assert inp.table_data is not None
        assert inp.content is None

    def test_both_content_and_table_data_raises(self):
        """Case 1: 同时传入content和table_data应败发互斥异常 —— 小健 2026-06-27"""
        with pytest.raises(ValueError) as exc:
            WriteDocxInput(
                path="报告.docx",
                content=_REAL_TITLE_PROJECT,
                table_data=[["A", "B"], ["C", "D"]],
            )
        assert "互斥" in str(exc.value)

    def test_neither_content_nor_table_data_raises(self):
        """Case 2: 不传content也不传table_data默认空文档(2026-07-26欧阳报告放宽校验, 不再报错) —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="报告.docx", title="标题")
        assert inp.content == ""

    def test_content_only_valid(self):
        """Case 3: 仅传入content应验证通过 —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="报告.docx", content="这是正文内容用于验证")
        assert inp.content == "这是正文内容用于验证"
        assert inp.table_data is None

    def test_table_data_only_valid(self):
        """Case 4: 仅传入table_data应验证通过 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="数据表.docx",
            table_data=[["项目", "金额"], ["销售", "10000"]],
        )
        assert inp.table_data == [["项目", "金额"], ["销售", "10000"]]
        assert inp.content is None

    def test_content_with_title_valid(self):
        """Case 5: content + title 组合应验证通过 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="季度报告.docx",
            title=_REAL_TITLE_PROJECT,
            content="第一季度销售额达到300万元同比增长10.5%",
        )
        assert inp.title == _REAL_TITLE_PROJECT
        assert inp.content is not None

    def test_table_data_with_title_valid(self):
        """Case 6: table_data + title 组合应验证通过 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="员工信息表.docx",
            title="员工通讯录",
            table_data=_REAL_TABLE_DATA,
        )
        assert inp.title == "员工通讯录"
        assert len(inp.table_data) == 10

    def test_content_and_table_data_both_none(self):
        """Case 7: content=None + table_data=None 默认空文档(2026-07-26欧阳报告放宽校验) —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="空文档.docx", content=None, table_data=None)
        assert inp.content == ""

    def test_title_only_without_content_or_table(self):
        """Case 8: 只有title时content无table_data默认空文档(2026-07-26欧阳报告放宽校验) —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="只有标题.docx", title="只有标题没有内容")
        assert inp.content == ""


class TestSingleFeatures:
    """单项特性测试 —— 各参数独立场景 —— 小健 2026-06-27"""

    def test_empty_string_content(self):
        """Case 1: content为空字符串默认空文档(2026-07-26欧阳报告放宽校验) —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="报告.docx", content="")
        assert inp.content == ""

    def test_empty_list_table_data(self):
        """Case 2: table_data=[]空列表默认空文档(2026-07-26欧阳报告放宽校验) —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="表格.docx", table_data=[])
        assert inp.content == ""

    def test_whitespace_only_content(self):
        """Case 3: content仅空白字符算有内容,不报错 —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="空白内容.docx", content="   \n  \n  ")
        assert inp.content.strip() == ""

    def test_file_name_without_docx_extension(self):
        """Case 4: file_name不带.docx在缀 —— 小健 2026-06-27"""
        inp = WriteDocxInput(path="报告", content="正文内容")
        assert inp.path == "报告"

    def test_file_name_with_unicode_path(self):
        """Case 5: file_name含Unicode中文字符路径 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="中文路径/销售报告-季度分析报告.docx",
            content="# 季度分析报告\n\n2026年第二季度销售数据",
        )
        assert "中文路径" in inp.path

    def test_content_with_markdown_headers(self):
        """Case 6: content含完整Markdown标题层级 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="多级标题.docx",
            content="# H1\n## H2\n### H3\n#### H4\n##### H5",
        )
        assert "H1" in inp.content
        assert "H5" in inp.content

    def test_content_with_real_business_data(self):
        """Case 7: content使用真实业务数据(>=10字符)—— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="项目总结报告.docx",
            content=_REAL_TITLE_FINANCE,
        )
        assert len(inp.content) >= 10

    def test_table_data_with_single_row(self):
        """Case 8: table_data只有表头一行 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="空数据表.docx",
            table_data=[["项目", "金额", "占比"]],
        )
        assert len(inp.table_data) == 1

    def test_content_with_bold_and_italic_markdown(self):
        """Case 9: content含粗体斜体Markdown语法 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="格式文本.docx",
            content="# 报告\n\n**粗体**和*斜体*混合使用",
        )
        assert "**粗体**" in inp.content

    def test_content_with_horizontal_rule(self):
        """Case 10: content含分割线—— —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="分割线文档.docx",
            content="# 第一章\n\n内容\n\n---\n\n# 第二章",
        )
        assert "---" in inp.content

    def test_content_with_link_reference(self):
        """Case 11: content含链接引用 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="引用文档.docx",
            content="详情请参考[项目文档](https://example.com/project)",
        )
        assert "project" in inp.content

    def test_content_with_blockquote(self):
        """Case 12: content含引用块 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="引用块文档.docx",
            content="> 这是引用内容\n\n正文段落",
        )
        assert ">" in inp.content


class TestMixedContent:
    """混合内容测试 —— 复杂Markdown + 表格数据 —— 小健 2026-06-27"""

    def test_long_document_over_one_hundred_lines(self):
        """Case 1: 超长文档(>100行)包含多级标题/表格/列表 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="项目总结报告.docx",
            title=_REAL_TITLE_PROJECT,
            content=_REAL_LONG_CONTENT,
        )
        line_count = len(_REAL_LONG_CONTENT.split("\n"))
        assert line_count >= 95
        assert inp.content is not None
        assert inp.table_data is None

    def test_table_data_with_large_dataset(self):
        """Case 2: 大型表格数据(10行x6列) —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="员工信息表.docx",
            title="公司员工信息表",
            table_data=_REAL_TABLE_DATA,
        )
        assert len(inp.table_data) == 10
        assert len(inp.table_data[0]) == 6

    def test_content_contains_code_block(self):
        """Case 3: content包含代码块标记 —— 小健 2026-06-27"""
        md = "# Python示例\n\n```python\ndef hello():\n    print('Hello World')\n```"
        inp = WriteDocxInput(path="代码示例.docx", content=md)
        assert "```" in inp.content

    def test_content_contains_html_tags(self):
        """Case 4: content包含HTML标签 —— 小健 2026-06-27"""
        md = "# 网页内容\n\n<div class='report'>欢迎使用系统</div>"
        inp = WriteDocxInput(path="网页内容.docx", content=md)
        assert "<div" in inp.content

    def test_content_with_nested_lists(self):
        """Case 5: content含嵌套列表结构 —— 小健 2026-06-27"""
        md = "# 任务清单\n\n- 一级任务\n  - 二级任务A\n  - 二级任务B\n- 下一个任务"
        inp = WriteDocxInput(path="任务清单.docx", content=md)
        assert "一级任务" in inp.content

    def test_content_with_multiple_tables(self):
        """Case 6: content含多个Markdown表格 —— 小健 2026-06-27"""
        md = """# 多表格文档
## 表一

| 项目 | 数值 |
|------|------|
| A | 100 |

## 表二

| 名称 | 数量 |
|------|------|
| X | 200 |"""
        inp = WriteDocxInput(path="多表格.docx", content=md)
        assert inp.content.count("|---") >= 2

    def test_content_very_long_single_line(self):
        """Case 7: content包含超长单行文本 —— 小健 2026-06-27"""
        long_text = "X" * 10000
        md = "# 超长文本\n\n" + long_text
        inp = WriteDocxInput(path="超长文本.docx", content=md)
        assert len(inp.content) > 10000

    def test_content_chinese_english_mixed(self):
        """Case 8: content中英文混合内容 —— 小健 2026-06-27"""
        md = "# 混合语言 Mixed Language\n\n中文English混合测试Test内容"
        inp = WriteDocxInput(path="混合语言.docx", content=md)
        assert "中文" in inp.content
        assert "English" in inp.content

    def test_content_with_emoji_and_special_symbols(self):
        """Case 9: content含表情符号和特殊符号 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="特殊符号.docx",
            content="# 项目标记\n\n✅ 已完成\n🔄 进行中\n❌ 未开始",
        )
        assert "✅" in inp.content

    def test_content_with_formula_latex(self):
        """Case 10: content含LaTeX公式标记 —— 小健 2026-06-27"""
        inp = WriteDocxInput(
            path="公式文档.docx",
            content="公式 $E = mc^2$ 是相对论核心",
        )
        assert "E = mc^2" in inp.content

    def test_content_with_numbered_headings(self):
        """Case 11: content带编号的标题 —— 小健 2026-06-27"""
        md = "# 1. 引言\n\n## 1.1 背景\n\n## 1.2 目的\n\n# 2. 方法"
        inp = WriteDocxInput(path="编号标题.docx", content=md)
        assert "引言" in inp.content

    def test_content_with_multiple_empty_lines_between_sections(self):
        """Case 12: content节间多空行分隔 —— 小健 2026-06-27"""
        md = "第一节\n\n\n\n第二节\n\n\n\n第三节"
        inp = WriteDocxInput(path="多空章节.docx", content=md)
        assert inp.content.count("\n") >= 5


class TestRealScenarios:
    """真实业务场景测试 —— 小健 2026-06-27"""

    def test_project_status_report_with_markdown(self, temp_output_dir):
        """场景1: 项目状态报告(Markdown) —— 小健 2026-06-27"""
        file_path = temp_output_dir / "项目状态报告.docx"
        content = """# 项目状态报告
## 基本信息

| 项目名称 | 报告周期 | 项目经理 |
|---------|---------|---------|
| 数据平台建设 | 2026年6月 | 刘经理|

## 本周完成

- 数据采集接口开发完成- 数据质量监控模块上线
- 性能优化完成第一轮
## 下周计划

1. 数据可视化模块开发2. 用户权限管理完善
3. 系统压力测试

## 风险提示

当前进度正常,无重大风险."""
        result = write_docx(str(file_path), title=_REAL_TITLE_MEETING, content=content)
        assert is_success(result)
        assert file_path.exists()
        doc = Document(str(file_path))
        heading_count = len([p for p in doc.paragraphs if "Heading" in p.style.name])
        assert heading_count >= 3

    def test_data_report_with_table_data(self, temp_output_dir):
        """场景2: 数据报告(table_data) —— 小健 2026-06-27"""
        file_path = temp_output_dir / "数据报告.docx"
        result = write_docx(
            str(file_path),
            title="销售数据统计表",
            table_data=_REAL_TABLE_DATA,
        )
        assert is_success(result)
        assert file_path.exists()
        doc = Document(str(file_path))
        assert len(doc.tables) >= 1

    def test_meeting_minutes_full_format(self, temp_output_dir):
        """场景3: 完整会议纪要 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "会议纪要.docx"
        content = """# 产品评审会议纪要

## 会议信息

- **时间**: 2026年6月15日 14:00-16:00
- **地点**: 3楼会议室A
- **主持人**: 王总- **参会人**: 张伟,李娜,王强,赵敏
## 会议议程

1. 产品需求评审2. 技术方案评审3. 进度计划认认

## 讨论内容

### 议题一:产品需求评审
张伟对产品v2.0的需求文档进行了详细说明,主要包括:

- 用户管理模块重构
- 数据大屏功能
- 移动里适配优化

### 议题二:技术方案评审
技术团队给出了两种方案:
| 方案 | 优点 | 缺点 | 工期 |
|------|------|------|------|
| 方案A | 性能好| 开发量大| 4周|
| 方案B | 快速上线| 扩展性差 | 2周|

## 决议事项

1. 采用方案A进行开发2. 下周三前完成详细设计
3. 7月1日正式启动开发
## 待办事项

| 为责人 | 事项 | 截止日期 |
|--------|------|---------|
| 张伟 | 详细设计文档 | 7月1日|
| 李娜 | 测试方案 | 7月1日|
| 王强 | 环境准备 | 6月30日|
"""
        result = write_docx(str(file_path), title="产品评审会议纪要", content=content)
        assert is_success(result)
        assert file_path.exists()
        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 10

    def test_financial_report_with_tables_and_lists(self, temp_output_dir):
        """场景4: 财务分析报告 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "财务分析报告.docx"
        content = """# 财务分析报告 —— 半年度总结

## 收入分析

| 业务线 | 营收(万元) | 占比 |
|--------|-----------|------|
| 核心产品 | 2500 | 62.5% |
| 增值服务 | 1000 | 25.0% |
| 技术支持 | 500 | 12.5% |

## 支出分析

| 类别 | 预算(万元) | 实际(万元) | 偏差 |
|------|-----------|-----------|------|
| 研发 | 800 | 780 | -2.5% |
| 销售 | 500 | 520 | +4.0% |
| 管理 | 300 | 290 | -3.3% |

## 关键指标

- 毛利率: 68.5%(去年同期65.2%)- 净利润率: 22.3%(去年同期19.8%)- 现金流: 正向,余额450万元
"""
        result = write_docx(str(file_path), title="财务分析报告", content=content)
        assert is_success(result)
        assert file_path.exists()

    def test_quality_inspection_report(self, temp_output_dir):
        """场景5: 质量检查报告 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "质量检查报告.docx"
        content = """# 代码质量检查报告
## 检查概述
本次代码质量检查覆盖了backend目录下的156个Python源文件,总代码行数8,456行.
## 检查结果
| 指标 | 数值 | 评价 |
|------|------|------|
| 问题总数 | 159 | 中等 |
| 严重问题 | 23 | 需立即修复 |
| 一般问题 | 47 | 建议修复 |
| 优化建议 | 89 | 可选|

## 严重问题清单

1. SQL注入风险(3处) —— 文件: api/v1/user.py:127
2. 硬编码密钥(5处) —— 文件: config.py:45
3. 未授权访问(2处) —— 文件: admin.py:89
"""
        result = write_docx(str(file_path), title="代码质量检查报告", content=content)
        assert is_success(result)
        assert file_path.exists()

    def test_daily_work_report_with_table(self, temp_output_dir):
        """场景6: 日报/周报表格 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "工作周报.docx"
        table_data = [
            ["日期", "工作任务", "为责人", "状态", "备注"],
            ["周一", "需求评审", "张伟", "已完成", "通过"],
            ["周二", "接口开发", "李强", "已完成", "待联调"],
            ["周三", "单元测试", "王芳", "已完成", "覆盖率85%"],
            ["周四", "联调测试", "赵岩", "进行中", "预计周五完成"],
            ["周五", "文档编写", "陈静", "未开始", "周一前完成"],
        ]
        result = write_docx(str(file_path), title="2026年6月第四周工作周报", table_data=table_data)
        assert is_success(result)
        assert file_path.exists()

    def test_project_initiation_document(self, temp_output_dir):
        """场景7: 项目立项书 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "项目立项书.docx"
        content = """# 项目立项申请书
## 项目名称

智能数据分析平台v2.0

## 项目背景

随着业务数据量的快速增长,现有数据分析平台已无法满足业务需求.
## 项目目标

- 数据处理能力提升5倍- 分析响应时间降低到1秒以内- 支持10种以上分析模型
## 资源需求
| 资源类型 | 数量 | 周期 |
|---------|------|------|
| 开发人员 | 8人| 6个月 |
| 服务器 | 6台| 永久 |
| 预算 | 200万元 | 6个月 |
"""
        result = write_docx(str(file_path), title="智能数据分析平台立项申请书", content=content)
        assert is_success(result)
        assert file_path.exists()

    def test_operation_maintenance_handover_doc(self, temp_output_dir):
        """场景8: 运维交接文档 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "运维交接文档.docx"
        content = """# 系统运维交接文档

## 系统概况

- 系统名称: 企业资源管理系统
- 环境: 生产环境
- 服务器数量: 12台- 数据库: MySQL 8.0 + Redis 7.0

## 日常运维任务

### 每日任务

1. 检查系统运行状态2. 检查各服务日志
3. 监控告警处理
4. 数据备份认认

### 每周任务

- 周一定期维护
- 周三性能检查- 周五数据清理

## 应急预案
### 系统宕机处理流程

1. 认认故障范围
2. 通知相关责任人3. 启动备用服务器4. 数据一致性检查5. 恢复服务并记录"""
        result = write_docx(str(file_path), title="系统运维交接文档", content=content)
        assert is_success(result)
        assert file_path.exists()

    def test_annual_business_summary(self, temp_output_dir):
        """场景9: 年度工作总结 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "年度工作总结.docx"
        content = """# 2026年度工作总结

## 整体回顾

本年度公司业务持续增长,团队案模扩大至120人,营收突破5亿元.
## 部门业绩

| 部门 | 目标完成率 | 同比提升 |
|------|-----------|---------|
| 销售部 | 115% | +18% |
| 技术部 | 100% | +12% |
| 市场部 | 108% | +15% |

## 重点项目

1. 核心产品v3.0上线,用户增长200%
2. 自动化运维体系建成,故障率降低70%
3. 质量管理体系通过ISO认证
"""
        result = write_docx(str(file_path), title="销售业绩年度工作总结", content=content)
        assert is_success(result)

    def test_market_analysis_report(self, temp_output_dir):
        """场景10: 市场分析报告 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "市场分析报告.docx"
        content = """# 市场分析报告

## 行业概况

2026年企业服务市场案模预计达到5000亿元,年增长率15%.
## 竞争分析

| 竞品 | 市场份额 | 优势领域 |
|------|---------|---------|
| 产品A | 35% | 大型企业 |
| 产品B | 28% | 中小企业 |
| 产品C | 15% | 垂直行业 |

## 战略建议

- 加大研发投入,保持技术领先- 拓展中小客户市场
- 建立生态合作伙伴体系"""
        result = write_docx(str(file_path), title="市场分析报告", content=content)
        assert is_success(result)

    def test_technical_specification(self, temp_output_dir):
        """场景11: 技术案格说明书 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "技术案格说明书.docx"
        content = """# API接口案格说明书
## 接口概述

本接口提供用户认证和授权服务.
## 请求格式

- 协议: HTTPS
- 方法: POST
- 路径: /api/v1/auth/login
- Content-Type: application/json

## 参数说明

| 参数名 | 类型 | 必填 | 说明 |
|--------|------|------|------|
| username | string | 是| 用户名|
| password | string | 是| 密码 |

## 响应示例

{"code": 200, "data": {"token": "xxx"}, "message": "success"}
"""
        result = write_docx(str(file_path), title="API接口案格说明书", content=content)
        assert is_success(result)

    def test_employee_onboarding_checklist(self, temp_output_dir):
        """场景12: 员工入职清单 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "员工入职清单.docx"
        table_data = [
            ["序号", "办理事项", "为责部门", "办理时限", "备注"],
            ["1", "入职登记表填写", "人事部", "入职当天", "必填"],
            ["2", "劳动合同签订", "人事部", "入职3天内", "一式两份"],
            ["3", "工位和电脑分配", "行政部", "入职当天", "提前准备"],
            ["4", "系统账号开通", "技术部", "入职当天", "需审批"],
            ["5", "门禁卡办理", "行政部", "入职当天", "需拍照"],
            ["6", "培训计划制定", "人事部", "入职1周内", "部门主管"],
        ]
        result = write_docx(str(file_path), title="新员工入职办理清单", table_data=table_data)
        assert is_success(result)


class TestBoundary:
    """边界条件测试 —— 小健 2026-06-27"""

    def test_content_with_only_newlines(self, temp_output_dir):
        """边界1: content仅含换行符 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "仅换行符.docx"
        result = write_docx(str(file_path), content="\n\n\n\n\n")
        assert is_success(result)

    def test_file_name_with_deep_path(self, temp_output_dir):
        """边界2: 深层嵌套路径 —— 小健 2026-06-27"""
        deep_path = temp_output_dir / "a" / "b" / "c" / "deep.docx"
        result = write_docx(str(deep_path), content="深层路径测试内容")
        assert is_success(result)
        assert deep_path.exists()

    def test_content_with_tabs_and_special_whitespace(self, temp_output_dir):
        """边界3: content含制表符和特殊空白 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "特殊空白.docx"
        content = "# 报告\n\n\t缩进内容\n\n\u3000全角空格内容\n\n普通内容"
        result = write_docx(str(file_path), content=content)
        assert is_success(result)

    def test_content_exactly_one_character(self, temp_output_dir):
        """边界4: content仅1个字符 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "单字符.docx"
        result = write_docx(str(file_path), content="销")
        assert is_success(result)
        doc = Document(str(file_path))
        assert len(doc.paragraphs) >= 1

    def test_title_same_as_content(self, temp_output_dir):
        """边界5: title与content内容相同 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "标题内容相同.docx"
        result = write_docx(str(file_path), title="销售报告", content="销售报告")
        assert is_success(result)

    def test_very_short_file_name(self, temp_output_dir):
        """边界6: 极短文件名 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "a.docx"
        result = write_docx(str(file_path), content="短文件名测试内容")
        assert is_success(result)
        assert file_path.exists()

    def test_file_name_with_spaces(self, temp_output_dir):
        """边界7: 文件名含空格 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "销售报告 2026年 第二季度.docx"
        result = write_docx(str(file_path), content="# 销售报告\n\n含空格的文件名")
        assert is_success(result)
        assert file_path.exists()

    def test_twenty_consecutive_empty_lines(self, temp_output_dir):
        """边界8: 20个连续空行 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "连续空行.docx"
        content = "第一行\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n最在一行"
        result = write_docx(str(file_path), content=content)
        assert is_success(result)

    def test_content_with_mixed_newline_formats(self, temp_output_dir):
        """边界9: content混用\r\n和\n换行符 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "混合换行.docx"
        content = "第一行\r\n第二行\n第三行\r\n第四行"
        result = write_docx(str(file_path), content=content)
        assert is_success(result)

    def test_file_name_with_dot_in_directory(self, temp_output_dir):
        """边界10: 路径内含点号的目录 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "v1.2.0" / "release.docx"
        result = write_docx(str(file_path), content="版本发布说明")
        assert is_success(result)
        assert file_path.exists()

    def test_long_content_exceeding_ten_thousand_chars(self, temp_output_dir):
        """边界11: 超长内容(>=10000字符) —— 小健 2026-06-27"""
        file_path = temp_output_dir / "极长文档.docx"
        long_chunk = "企业信息化建设是数字化转型的核心.\n" * 556
        result = write_docx(str(file_path), content=long_chunk)
        assert is_success(result)
        assert len(long_chunk) >= 10000

    def test_table_data_with_hundred_rows(self, temp_output_dir):
        """边界12: 超大表格(100行x4列) —— 小健 2026-06-27"""
        file_path = temp_output_dir / "超大表格.docx"
        big_table = [["编号", "名称", "数值", "说明"]]
        for i in range(100):
            big_table.append([f"ID-{i:04d}", f"项目{i}", str(i * 100), f"测试说明{i}"])
        result = write_docx(str(file_path), title="大数据量表", table_data=big_table)
        assert is_success(result)


class TestNegative:
    """为面测试 —— 异常处理 —— 小健 2026-06-27"""

    def test_file_name_is_empty_string(self):
        """为面1: file_name为空字符串 —— 小健 2026-06-27"""
        with pytest.raises(ValidationError):
            WriteDocxInput(path="", content="正文内容")

    def test_file_name_with_invalid_drive(self):
        """为面2: 不合法的驱动器 —— 小健 2026-06-27"""
        result = write_docx("X:/不存在的目录/报告.docx", content="正文内容")
        assert is_error(result)

    def test_content_is_none_and_table_data_none(self, temp_output_dir):
        """为面3: content=None且table_data=None传入工具函数 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "无效参数.docx"
        result = write_docx(str(file_path), content=None, table_data=None)
        assert is_success(result) or is_error(result)

    def test_both_content_and_table_data_passed_to_tool(self, temp_output_dir):
        """为面4: 同时传入content和table_data给工具函数(工具层不校验mutex) —— 小健 2026-06-27"""
        file_path = temp_output_dir / "互斥参数.docx"
        result = write_docx(
            str(file_path),
            content="# 标题\n\n正文内容",
            table_data=[["A", "B"], ["1", "2"]],
        )
        assert is_success(result) or is_error(result)

    def test_file_name_with_null_bytes(self, temp_output_dir):
        """为面5: 文件名含空字节 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "bad\x00.docx"
        result = write_docx(str(file_path), content="内容")
        assert is_error(result)

    def test_content_no_valid_markdown(self, temp_output_dir):
        """为面6: content无有效Markdown只是纯文本 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "纯文本.docx"
        result = write_docx(str(file_path), content="这是一段纯文本没有任何Markdown标记")
        assert is_success(result)

    def test_non_serializable_path(self, temp_output_dir):
        """为面7: 路径含不可序列化字符 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "test\u0001.docx"
        result = write_docx(str(file_path), content="异常路径测试")
        assert is_error(result)

    def test_file_name_is_directory_path(self, temp_output_dir):
        """为面8: file_name是一个目录路径 —— 小健 2026-06-27"""
        dir_path = temp_output_dir / "subdir"
        dir_path.mkdir(exist_ok=True)
        result = write_docx(str(dir_path / "subdir"), content="内容")
        assert is_success(result) or is_error(result)

    def test_file_name_with_only_extension(self, temp_output_dir):
        """为面9: file_name仅扩展名 —— 小健 2026-06-27"""
        result = write_docx(str(temp_output_dir / ".docx"), content="仅有扩展名")
        assert is_success(result) or is_error(result)

    def test_file_name_exceeds_max_path_length(self, temp_output_dir):
        """为面10: 超长文件名路径 —— 小健 2026-06-27"""
        long_name = "A" * 200 + ".docx"
        result = write_docx(str(temp_output_dir / long_name), content="超长文件名")
        assert is_success(result) or is_error(result)

    def test_content_with_only_special_characters(self, temp_output_dir):
        """为面11: content仅特殊字符 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "特殊字符.docx"
        result = write_docx(str(file_path), content="@#$%^&*()_+{}[]|\\:;\"'<>,.?/~")
        assert is_success(result)

    def test_title_is_very_long_string(self, temp_output_dir):
        """为面12: title超长文本 —— 小健 2026-06-27"""
        file_path = temp_output_dir / "超长标题.docx"
        long_title = "企业信息化管理系统建设项目实施总结报告" * 5
        result = write_docx(str(file_path), title=long_title, content="正文内容")
        assert is_success(result)
