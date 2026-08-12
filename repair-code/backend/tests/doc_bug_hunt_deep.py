# -*- coding: utf-8 -*-
"""
Document Tools Bug Hunt - Deep Dive Phase 2 —小健 2026-06-24

Targeted tests for suspicious behaviors found in Phase 1.
Directly calls tool functions with edge cases to find real bugs.

Usage:
    cd G:\OmniAgentAs-desk\backend
    python tests\doc_bug_hunt_deep.py
"""

import sys
import os
import tempfile
import traceback
import yaml
import json

# Setup backend path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Create minimal config.yaml to bypass config initialization crash
_tmp_config_dir = tempfile.mkdtemp(prefix="bughunt_cfg_")
_tmp_config_path = os.path.join(_tmp_config_dir, "config.yaml")
with open(_tmp_config_path, "w", encoding="utf-8") as _f:
    yaml.dump({"app": {}, "ai": {}, "logging": {"level": "INFO"}}, _f)
os.environ["OMNIAGENT_CONFIG_PATH"] = _tmp_config_path

BUGS_FOUND = []
WARNINGS = []


def report_bug(test_name, description, details=""):
    BUGS_FOUND.append({"test": test_name, "description": description, "details": details})
    print(f"  [BUG] {description}")
    if details:
        for line in details.split("\n"):
            print(f"        {line}")


def report_warning(test_name, description, details=""):
    WARNINGS.append({"test": test_name, "description": description, "details": details})
    print(f"  [WARN] {description}")
    if details:
        for line in details.split("\n"):
            print(f"         {line}")


def report_ok(test_name, description=""):
    print(f"  [OK]   {description or test_name}")


def safe_call(func, *args, **kwargs):
    try:
        result = func(*args, **kwargs)
        return result, None
    except Exception as e:
        return None, e


def get_result_info(result):
    if result is None:
        return "None result"
    llm = result.get("llm_data", {})
    status = llm.get("status", {}) if llm else {}
    return status.get("exec_code", "unknown")


# ============================================================
# DEEP TEST 1: write_docx - empty table creation bug
# ============================================================
def test_write_docx_empty_table():
    print("\n" + "="*70)
    print("DEEP TEST: write_docx - empty table creation")
    print("="*70)
    from app.tools.document.write_docx import write_docx
    from docx import Document

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: table_data=[[]] creates empty table in document
    print("\n--- 1.1 table_data=[[]] - check for empty table in doc ---")
    path = os.path.join(tmpdir, "empty_tbl.docx")
    result = write_docx(path=path, table_data=[[]])
    if get_result_info(result) == "success":
        doc = Document(path)
        table_count = len(doc.tables)
        para_count = len(doc.paragraphs)
        print(f"    tables={table_count}, paragraphs={para_count}")
        if table_count > 0:
            # Check if the table is actually empty
            tbl = doc.tables[0]
            rows = len(tbl.rows)
            cols = len(tbl.columns)
            report_bug("write_docx_empty_tbl", 
                f"table_data=[[]] creates empty table ({rows}x{cols}) in document",
                "The function doesn't check if table_rows is empty before creating table.\n"
                f"  table_rows after parse = {rows} rows, {cols} cols\n"
                "  This creates a corrupted/empty table in the docx file")
        else:
            report_ok("write_docx_empty_tbl", "No empty table created")

    # BUG CANDIDATE: table_data with all-empty rows
    print("\n--- 1.2 table_data with all empty cells ---")
    path = os.path.join(tmpdir, "all_empty_cells.docx")
    table_data = [["", "", ""], ["", "", ""]]
    result = write_docx(path=path, table_data=table_data)
    if get_result_info(result) == "success":
        doc = Document(path)
        table_count = len(doc.tables)
        print(f"    tables={table_count}")
        if table_count > 0:
            tbl = doc.tables[0]
            print(f"    table: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
            # Check cell content
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    if cell.text:
                        print(f"    cell[{ri}][{ci}] = '{cell.text}'")
        report_ok("write_docx_all_empty", f"Created {table_count} table(s)")

    # BUG CANDIDATE: Markdown table with only header row (no separator)
    print("\n--- 1.3 Markdown table with only header (no separator) ---")
    path = os.path.join(tmpdir, "header_only_table.docx")
    content = "# Title\n\n| Col1 | Col2 |\n"
    result = write_docx(path=path, content=content)
    if get_result_info(result) == "success":
        doc = Document(path)
        table_count = len(doc.tables)
        para_texts = [p.text for p in doc.paragraphs]
        print(f"    tables={table_count}, paragraphs={para_texts}")
        if table_count > 0:
            report_warning("write_docx_header_only_table",
                f"Header-only table creates {table_count} table(s) - may be unintended",
                "A markdown table without separator row (|---|---|) still creates a table")
        else:
            report_ok("write_docx_header_only_table", "No table created from header-only")

    # BUG CANDIDATE: Markdown table with separator only (no header)
    print("\n--- 1.4 Markdown table with separator only ---")
    path = os.path.join(tmpdir, "separator_only.docx")
    content = "# Title\n\n|------|------|\n| data1 | data2 |\n"
    result = write_docx(path=path, content=content)
    if get_result_info(result) == "success":
        doc = Document(path)
        table_count = len(doc.tables)
        print(f"    tables={table_count}")
        if table_count > 0:
            tbl = doc.tables[0]
            print(f"    table: {len(tbl.rows)} rows")
            for ri, row in enumerate(tbl.rows):
                cells = [cell.text for cell in row.cells]
                print(f"    row[{ri}] = {cells}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 2: read_docx - tmp_path scoping bug
# ============================================================
def test_read_docx_tmp_path():
    print("\n" + "="*70)
    print("DEEP TEST: read_docx - tmp_path NameError")
    print("="*70)
    from app.tools.document.read_docx import read_docx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: tmp_path undefined when suffix != ".doc"
    # The finally block at line 126 does: os.unlink(tmp_path)
    # But tmp_path is only defined inside the if suffix == ".doc" block
    # If suffix != ".doc" and the code reaches the finally, NameError occurs

    # Test with a valid .docx file
    print("\n--- 2.1 Read .docx file (tmp_path should NOT be referenced) ---")
    from app.tools.document.write_docx import write_docx
    docx_path = os.path.join(tmpdir, "test.docx")
    write_docx(path=docx_path, content="test content")
    
    result, err = safe_call(read_docx, path=docx_path)
    if err:
        report_bug("read_docx_tmp_path", "Read .docx file crashes", str(err))
    else:
        code = get_result_info(result)
        report_ok("read_docx_tmp_path", f".docx read OK: exec_code={code}")

    # Test with non-existent .docx file (should hit error before finally, but check)
    print("\n--- 2.2 Read non-existent .docx file (finally block check) ---")
    result, err = safe_call(read_docx, path="C:\\nonexist\\test.docx")
    if err:
        report_bug("read_docx_tmp_path", "Non-existent .docx crashes in finally", str(err))
    else:
        code = get_result_info(result)
        report_ok("read_docx_tmp_path", f"Non-existent .docx: exec_code={code}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 3: read_xlsx CSV - sheet_names KeyError
# ============================================================
def test_read_xlsx_csv_keyerror():
    print("\n" + "="*70)
    print("DEEP TEST: read_xlsx - CSV sheet_names KeyError")
    print("="*70)
    from app.tools.document.read_xlsx import read_xlsx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: For CSV files, _read_csv_stdlib_inner returns 
    # {"headers": [...], "rows": [...], "row_count": N}
    # WITHOUT "sheet_names" key.
    # Then in read_xlsx line 214: sheet_count = len(result.get("sheet_names", []))
    # result.get("sheet_names", []) returns [] -> len() = 0 -> OK
    # Wait, .get("sheet_names", []) has a default of [] so it won't KeyError!
    # But the metrics will show sheet_count=0 for CSV, which may be misleading.
    # Let me verify.

    print("\n--- 3.1 CSV read - check llm_data metrics ---")
    csv_path = os.path.join(tmpdir, "test.csv")
    with open(csv_path, "w", encoding="utf-8", newline="") as f:
        f.write("Name,Age\nAlice,30\nBob,25\n")
    
    result, err = safe_call(read_xlsx, path=csv_path)
    if err:
        report_bug("read_xlsx_csv", "CSV read crashes", str(err))
    else:
        code = get_result_info(result)
        llm_data = result.get("llm_data", {})
        metrics = llm_data.get("metrics", {})
        sheet_count_metric = metrics.get("sheet_count", {})
        print(f"    exec_code={code}")
        print(f"    sheet_count metric = {sheet_count_metric}")
        if sheet_count_metric.get("value") == 0 and sheet_count_metric.get("text") == "0一〃":
            report_warning("read_xlsx_csv_sheet_count",
                "CSV read shows sheet_count=0 in metrics (misleading)",
                "The llm_data.summary and metrics show 0 sheets for CSV files\n"
                "  even though CSV is a valid single-sheet format.\n"
                "  This is confusing in the UI.")
        else:
            report_ok("read_xlsx_csv_sheet_count", f"sheet_count metric: {sheet_count_metric}")

    # Verify CSV data is actually correct
    print("\n--- 3.2 CSV data verification ---")
    data = result.get("data", {})
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    print(f"    headers={headers}")
    print(f"    rows={rows}")
    if headers == ["Name", "Age"] and len(rows) == 2:
        report_ok("read_xlsx_csv_data", "CSV data correct")
    else:
        report_bug("read_xlsx_csv_data", f"CSV data wrong: headers={headers}, rows={len(rows)}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 4: write_pptx - coerce_json with non-JSON string
# ============================================================
def test_write_pptx_coerce():
    print("\n" + "="*70)
    print("DEEP TEST: write_pptx - coerce_json with invalid input")
    print("="*70)
    from app.tools.document.write_pptx import write_pptx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: coerce_json("not json") returns "not json" as-is
    # Then _build_pptx_presentation iterates "not json" char by char
    # Each char goes to _add_pptx_slide which skips non-dict items
    # Result: 0 slides, success response - SILENT FAILURE

    print("\n--- 4.1 slides='not json' (invalid JSON string) ---")
    path = os.path.join(tmpdir, "bad_json.pptx")
    result, err = safe_call(write_pptx, path=path, slides="not json")
    if err:
        report_bug("write_pptx_coerce", "Invalid JSON string crashes", str(err))
    else:
        code = get_result_info(result)
        slide_count = result.get("data", {}).get("slide_count", 0)
        print(f"    exec_code={code}, slide_count={slide_count}")
        if code == "success" and slide_count == 0:
            report_bug("write_pptx_coerce",
                "Invalid JSON string silently creates empty presentation",
                "slides='not json' is not valid JSON.\n"
                "  coerce_json returns it as-is (a string).\n"
                "  _build_pptx_presentation iterates chars, each skipped by _add_pptx_slide.\n"
                "  User gets success with 0 slides instead of an error.\n"
                "  Expected: error response indicating invalid input")
        elif code == "success" and slide_count > 0:
            report_warning("write_pptx_coerce",
                f"Invalid JSON string creates {slide_count} slides from chars",
                "Each character is treated as a slide")
        else:
            report_ok("write_pptx_coerce", f"Handled: exec_code={code}")

    # Test with JSON array of non-dicts
    print("\n--- 4.2 slides='[1,2,3]' (JSON array of integers) ---")
    path = os.path.join(tmpdir, "json_ints.pptx")
    result, err = safe_call(write_pptx, path=path, slides="[1,2,3]")
    if err:
        report_bug("write_pptx_coerce", "JSON int array crashes", str(err))
    else:
        code = get_result_info(result)
        slide_count = result.get("data", {}).get("slide_count", 0)
        print(f"    exec_code={code}, slide_count={slide_count}")
        if code == "success" and slide_count == 0:
            report_bug("write_pptx_coerce",
                "JSON integer array silently creates empty presentation",
                "slides='[1,2,3]' parses to [1, 2, 3].\n"
                "  Each int fails isinstance(slide_data, dict) check.\n"
                "  All skipped. User gets 0 slides with success status.")
        else:
            report_ok("write_pptx_coerce", f"Handled: exec_code={code}")

    # Test with JSON object (not array)
    print("\n--- 4.3 slides='{}' (JSON object instead of array) ---")
    path = os.path.join(tmpdir, "json_obj.pptx")
    result, err = safe_call(write_pptx, path=path, slides="{}")
    if err:
        report_bug("write_pptx_coerce", "JSON object crashes", str(err))
    else:
        code = get_result_info(result)
        slide_count = result.get("data", {}).get("slide_count", 0)
        print(f"    exec_code={code}, slide_count={slide_count}")
        if code == "success" and slide_count == 0:
            report_warning("write_pptx_coerce",
                "JSON object {} silently creates empty presentation",
                "slides='{}' parses to {} (dict, not list).\n"
                "  _build_pptx_presentation iterates over dict keys (empty).\n"
                "  User gets 0 slides with success status.")
        elif code == "success":
            report_ok("write_pptx_coerce", f"Handled: exec_code={code}, slides={slide_count}")
        else:
            report_ok("write_pptx_coerce", f"Handled: exec_code={code}")

    # Test with valid JSON but list of strings
    print("\n--- 4.4 slides='[\"a\",\"b\"]' (JSON array of strings) ---")
    path = os.path.join(tmpdir, "json_strs.pptx")
    result, err = safe_call(write_pptx, path=path, slides='["a","b"]')
    if err:
        report_bug("write_pptx_coerce", "JSON string array crashes", str(err))
    else:
        code = get_result_info(result)
        slide_count = result.get("data", {}).get("slide_count", 0)
        print(f"    exec_code={code}, slide_count={slide_count}")
        if code == "success" and slide_count == 0:
            report_warning("write_pptx_coerce",
                "JSON string array silently creates empty presentation",
                "slides='[\"a\",\"b\"]' parses to ['a', 'b'].\n"
                "  Each string fails isinstance(slide_data, dict) check.\n"
                "  All skipped. User gets 0 slides with success status.")
        else:
            report_ok("write_pptx_coerce", f"Handled: exec_code={code}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 5: write_pdf - empty table_data error
# ============================================================
def test_write_pdf_empty_table():
    print("\n" + "="*70)
    print("DEEP TEST: write_pdf - table_data edge cases")
    print("="*70)
    from app.tools.document.write_pdf import write_pdf

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: table_data=[[]] passes truthy check, crashes in _create_pdf_table
    print("\n--- 5.1 table_data=[[]] ---")
    path = os.path.join(tmpdir, "empty_tbl.pdf")
    result, err = safe_call(write_pdf, path=path, table_data=[[]])
    if err:
        report_warning("write_pdf_empty_tbl",
            "table_data=[[]] crashes with exception",
            str(err))
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            print(f"    Error detail: {err_detail[:200]}")
            # Check if it's a meaningful error or just a crash
            if "index" in str(err_detail).lower() or "range" in str(err_detail).lower():
                report_bug("write_pdf_empty_tbl",
                    "table_data=[[]] causes IndexError/crash",
                    f"Error: {err_detail[:200]}\n"
                    "  table_data=[[]] is truthy, passes 'elif table_data:' check.\n"
                    "  _create_pdf_table receives [[]], tries to access table_data[0]=[]\n"
                    "  reportlab Table([]) may crash on empty rows.")
            else:
                report_warning("write_pdf_empty_tbl",
                    f"table_data=[[]] returns error: {err_detail[:100]}")
        else:
            report_ok("write_pdf_empty_tbl", "Empty table handled gracefully")

    # Test table_data with single empty row
    print("\n--- 5.2 table_data=[['','']] ---")
    path = os.path.join(tmpdir, "empty_str_rows.pdf")
    result, err = safe_call(write_pdf, path=path, table_data=[["", ""]])
    if err:
        report_bug("write_pdf_empty_str_rows", "table_data=[['','']] crashes", str(err))
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_pdf_empty_str_rows",
                f"table_data=[['','']] returns error: {err_detail[:100]}")
        else:
            report_ok("write_pdf_empty_str_rows", "Empty string rows handled")

    # Test table_data with uneven rows
    print("\n--- 5.3 table_data with uneven rows ---")
    path = os.path.join(tmpdir, "uneven_rows.pdf")
    result, err = safe_call(write_pdf, path=path, table_data=[["A", "B", "C"], ["1"]])
    if err:
        report_warning("write_pdf_uneven",
            "Uneven table_data crashes",
            str(err))
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_pdf_uneven",
                f"Uneven rows error: {err_detail[:100]}")
        else:
            report_ok("write_pdf_uneven", "Uneven rows handled")

    # Test content with malformed markdown table
    print("\n--- 5.4 Content with malformed table ---")
    path = os.path.join(tmpdir, "bad_table.pdf")
    content = "# Test\n\n| A | B |\nno separator here\n| C | D |\n"
    result, err = safe_call(write_pdf, path=path, content=content)
    if err:
        report_bug("write_pdf_bad_table", "Malformed table crashes", str(err))
    else:
        code = get_result_info(result)
        report_ok("write_pdf_bad_table", f"Malformed table handled: exec_code={code}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 6: write_xlsx - edge cases for data parameter
# ============================================================
def test_write_xlsx_data():
    print("\n" + "="*70)
    print("DEEP TEST: write_xlsx - data parameter edge cases")
    print("="*70)
    from app.tools.document.write_xlsx import write_xlsx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: coerce_json("not json") returns "not json" as string
    # Then data is a string, len(data) > 0 is True
    # headers = list(dict.fromkeys(k for row in data for k in row.keys()))
    # This iterates over characters of string, calling .keys() on each char -> AttributeError
    print("\n--- 6.1 data='not json' (invalid JSON string) ---")
    path = os.path.join(tmpdir, "bad_str.xlsx")
    result, err = safe_call(write_xlsx, path=path, data="not json")
    if err:
        report_bug("write_xlsx_bad_str",
            "Invalid JSON string crashes with AttributeError",
            f"Exception: {type(err).__name__}: {err}\n"
            "  coerce_json('not json') returns 'not json' (string)\n"
            "  data='not json' passes 'if len(data) > 0' check\n"
            "  'for row in data' iterates chars\n"
            "  'row.keys()' on a char raises AttributeError")
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_xlsx_bad_str",
                f"Invalid string handled with error: {err_detail[:100]}")
        else:
            # If it returned success, check what happened
            print(f"    exec_code={code}")
            report_ok("write_xlsx_bad_str", f"Handled: exec_code={code}")

    # BUG CANDIDATE: data as integer
    print("\n--- 6.2 data=12345 (integer) ---")
    path = os.path.join(tmpdir, "int_data.xlsx")
    result, err = safe_call(write_xlsx, path=path, data=12345)
    if err:
        report_bug("write_xlsx_int_data",
            "Integer data crashes",
            f"Exception: {type(err).__name__}: {err}")
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_xlsx_int_data",
                f"Integer data handled with error: {err_detail[:100]}")
        else:
            report_ok("write_xlsx_int_data", f"Integer data handled: exec_code={code}")

    # BUG CANDIDATE: data as dict (not list of dicts)
    print("\n--- 6.3 data={'a': 1} (dict instead of list) ---")
    path = os.path.join(tmpdir, "dict_data.xlsx")
    result, err = safe_call(write_xlsx, path=path, data={"a": 1})
    if err:
        report_warning("write_xlsx_dict_data",
            "Dict data crashes",
            f"Exception: {type(err).__name__}: {err}")
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_xlsx_dict_data",
                f"Dict data error: {err_detail[:100]}")
        else:
            report_ok("write_xlsx_dict_data", f"Dict data handled: exec_code={code}")

    # Edge case: data with None values
    print("\n--- 6.4 data with None values in cells ---")
    path = os.path.join(tmpdir, "none_vals.xlsx")
    data = [{"A": None, "B": "text"}, {"A": 123, "B": None}]
    result, err = safe_call(write_xlsx, path=path, data=data)
    if err:
        report_bug("write_xlsx_none_vals", "None values crash", str(err))
    else:
        code = get_result_info(result)
        report_ok("write_xlsx_none_vals", f"None values handled: exec_code={code}")

    # Edge case: nested dict value
    print("\n--- 6.5 data with nested dict value ---")
    path = os.path.join(tmpdir, "nested.xlsx")
    data = [{"A": {"nested": "value"}, "B": "text"}]
    result, err = safe_call(write_xlsx, path=path, data=data)
    if err:
        report_warning("write_xlsx_nested",
            "Nested dict value crashes",
            str(err))
    else:
        code = get_result_info(result)
        if code == "error":
            err_detail = result.get("llm_data", {}).get("status", {}).get("detail", "")
            report_warning("write_xlsx_nested",
                f"Nested dict error: {err_detail[:100]}")
        else:
            report_ok("write_xlsx_nested", f"Nested dict handled: exec_code={code}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 7: write_pptx - _select_layout KeyError
# ============================================================
def test_write_pptx_layout():
    print("\n" + "="*70)
    print("DEEP TEST: write_pptx - _select_layout with various type values")
    print("="*70)
    from app.tools.document.write_pptx import write_pptx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # _select_layout maps: {0: 0, "cover": 0, 1: 1, "content": 1, 2: 2, "two": 2}
    # .get(slide_type, 1) defaults to 1 for unknown types

    # Test with integer type values
    for type_val in [0, 1, 2, 3, -1, 99]:
        print(f"\n--- 7.{type_val+10} slide type={type_val} ---")
        path = os.path.join(tmpdir, f"type_{type_val}.pptx")
        slides = [{"title": f"Type {type_val}", "type": type_val}]
        result, err = safe_call(write_pptx, path=path, slides=slides)
        if err:
            report_bug("write_pptx_layout", f"type={type_val} crashes", str(err))
        else:
            code = get_result_info(result)
            slide_count = result.get("data", {}).get("slide_count", 0)
            print(f"    exec_code={code}, slide_count={slide_count}")

    # Test with string type values
    for type_val in ["cover", "content", "two", "invalid", "", "COVER"]:
        print(f"\n--- 7.x slide type='{type_val}' ---")
        path = os.path.join(tmpdir, f"type_{type_val or 'empty'}.pptx")
        slides = [{"title": f"Type {type_val}", "type": type_val}]
        result, err = safe_call(write_pptx, path=path, slides=slides)
        if err:
            report_bug("write_pptx_layout", f"type='{type_val}' crashes", str(err))
        else:
            code = get_result_info(result)
            slide_count = result.get("data", {}).get("slide_count", 0)
            print(f"    exec_code={code}, slide_count={slide_count}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 8: Round-trip consistency checks
# ============================================================
def test_roundtrip_consistency():
    print("\n" + "="*70)
    print("DEEP TEST: Round-trip consistency")
    print("="*70)
    from app.tools.document.write_docx import write_docx
    from app.tools.document.read_docx import read_docx
    from app.tools.document.write_xlsx import write_xlsx
    from app.tools.document.read_xlsx import read_xlsx
    from app.tools.document.write_pdf import write_pdf
    from app.tools.document.read_pdf import read_pdf
    from app.tools.document.write_pptx import write_pptx
    from app.tools.document.read_pptx import read_pptx

    tmpdir = tempfile.mkdtemp(prefix="bughunt_rt_")

    # 8.1: DOCX round-trip with markdown content
    print("\n--- 8.1 DOCX: Markdown content round-trip ---")
    docx_path = os.path.join(tmpdir, "rt.docx")
    original_md = "# Title\n\n## Subtitle\n\nParagraph with **bold** and *italic*.\n\n- Item 1\n- Item 2\n\n1. First\n2. Second\n"
    write_result = write_docx(path=docx_path, content=original_md)
    if get_result_info(write_result) == "success":
        read_result = read_docx(path=docx_path)
        if get_result_info(read_result) == "success":
            text = read_result.get("data", {}).get("text", "")
            # Check that headings survived
            has_title = "Title" in text
            has_subtitle = "Subtitle" in text
            has_items = "Item 1" in text
            print(f"    title={has_title}, subtitle={has_subtitle}, items={has_items}")
            if has_title and has_subtitle and has_items:
                report_ok("rt_docx_md", "Markdown round-trip OK")
            else:
                report_bug("rt_docx_md",
                    "Markdown content lost in round-trip",
                    f"Text: {text[:300]}")

    # 8.2: XLSX round-trip with mixed data types
    print("\n--- 8.2 XLSX: Mixed data types round-trip ---")
    xlsx_path = os.path.join(tmpdir, "rt.xlsx")
    original_data = [
        {"Name": "Alice", "Age": 30, "Score": 95.5, "Active": True},
        {"Name": "Bob", "Age": 25, "Score": None, "Active": False},
    ]
    write_result = write_xlsx(path=xlsx_path, data=original_data)
    if get_result_info(write_result) == "success":
        read_result = read_xlsx(path=xlsx_path)
        if get_result_info(read_result) == "success":
            data = read_result.get("data", {})
            headers = data.get("headers", [])
            rows = data.get("rows", [])
            print(f"    headers={headers}")
            for i, row in enumerate(rows):
                print(f"    row[{i}]={row}")
            # Check data preserved
            if rows and rows[0][0] == "Alice":
                report_ok("rt_xlsx_mixed", "Mixed data round-trip OK")
            else:
                report_bug("rt_xlsx_mixed", f"Data corrupted: {rows}")

    # 8.3: PPTX round-trip with table
    print("\n--- 8.3 PPTX: Table round-trip ---")
    pptx_path = os.path.join(tmpdir, "rt.pptx")
    slides = [
        {"title": "Table Slide", "tables": [[["A", "B"], ["1", "2"], ["3", "4"]]]},
    ]
    write_result = write_pptx(path=pptx_path, slides=slides)
    if get_result_info(write_result) == "success":
        read_result = read_pptx(path=pptx_path)
        if get_result_info(read_result) == "success":
            data = read_result.get("data", {})
            slides_data = data.get("slides", [])
            if slides_data and slides_data[0].get("tables"):
                tables = slides_data[0]["tables"]
                print(f"    tables found: {len(tables)}")
                if tables and tables[0] and tables[0][0] == ["A", "B"]:
                    report_ok("rt_pptx_table", "PPTX table round-trip OK")
                else:
                    report_bug("rt_pptx_table", f"Table data wrong: {tables}")
            else:
                report_warning("rt_pptx_table", "No tables found in read result")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 9: write_docx - content+table_data mutual exclusion
# ============================================================
def test_write_docx_exclusion():
    print("\n" + "="*70)
    print("DEEP TEST: write_docx - content vs table_data exclusion")
    print("="*70)
    from app.tools.document.write_docx import write_docx
    from docx import Document

    tmpdir = tempfile.mkdtemp(prefix="bughunt_deep_")

    # BUG CANDIDATE: When both content and table_data are provided,
    # content is used and table_data is ignored (elif branch).
    # But there's no warning to the user about this.
    print("\n--- 9.1 Both content and table_data provided ---")
    path = os.path.join(tmpdir, "both.docx")
    content = "# Content Section\n\nThis is content."
    table_data = [["X", "Y"], ["1", "2"]]
    result = write_docx(path=path, content=content, table_data=table_data)
    if get_result_info(result) == "success":
        doc = Document(path)
        table_count = len(doc.tables)
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        print(f"    tables={table_count}, paragraphs={para_texts}")
        if table_count > 0:
            report_warning("write_docx_exclusion",
                "Both content+table_data: table_data was NOT ignored",
                "Schema says they're mutually exclusive and content takes priority,\n"
                "  but both content and table_data were written to the document.")
        elif "Content Section" in str(para_texts):
            report_ok("write_docx_exclusion", "Content took priority, table_data ignored (correct)")
        else:
            report_warning("write_docx_exclusion", f"Unexpected: tables={table_count}, paras={para_texts}")

    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


# ============================================================
# DEEP TEST 10: read_docx - .doc file temp file cleanup
# ============================================================
def test_read_docx_doc_cleanup():
    print("\n" + "="*70)
    print("DEEP TEST: read_docx - .doc temp file cleanup")
    print("="*70)
    from app.tools.document.read_docx import read_docx

    # BUG CANDIDATE: In read_docx, when suffix == ".doc":
    # - tmp_path is defined at line 66
    # - If subprocess.run fails, the except block returns error
    # - But the finally block at line 76 does: pass
    # - The actual cleanup is at line 126-129 in the outer finally
    # - If pandoc conversion fails, the function returns from line 74
    # - The outer finally block (126-129) tries to os.unlink(tmp_path)
    # - But if the error was at line 63 (pandoc not found), tmp_path IS defined
    # - So this should work. But let me verify the flow.

    # For .doc files without pandoc:
    print("\n--- 10.1 .doc file without pandoc ---")
    result, err = safe_call(read_docx, path="C:\\test.doc")
    if err:
        # This is expected - pandoc not found
        report_ok("read_docx_doc_pandoc", f".doc without pandoc: {type(err).__name__}")
    else:
        code = get_result_info(result)
        if code == "error":
            report_ok("read_docx_doc_pandoc", "Correctly returns error for missing pandoc")
        else:
            report_ok("read_docx_doc_pandoc", f"Handled: exec_code={code}")

    # For .doc file with invalid path (pandoc exists but file doesn't):
    print("\n--- 10.2 .doc file with invalid path (pandoc exists) ---")
    import shutil
    pandoc = shutil.which("pandoc")
    if pandoc:
        result, err = safe_call(read_docx, path="C:\\nonexist\\test.doc")
        if err:
            report_ok("read_docx_doc_nofile", f"Invalid .doc path: {type(err).__name__}")
        else:
            code = get_result_info(result)
            report_ok("read_docx_doc_nofile", f"Handled: exec_code={code}")
    else:
        report_ok("read_docx_doc_nofile", "Pandoc not installed, skipping")


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("=" * 70)
    print("DOCUMENT TOOLS BUG HUNT - DEEP DIVE PHASE 2")
    print("=" * 70)
    print("Running deep analysis...\n")

    try:
        test_write_docx_empty_table()
    except Exception as e:
        print(f"\n[FATAL] write_docx empty table test crashed: {e}")
        traceback.print_exc()

    try:
        test_read_docx_tmp_path()
    except Exception as e:
        print(f"\n[FATAL] read_docx tmp_path test crashed: {e}")
        traceback.print_exc()

    try:
        test_read_xlsx_csv_keyerror()
    except Exception as e:
        print(f"\n[FATAL] read_xlsx CSV test crashed: {e}")
        traceback.print_exc()

    try:
        test_write_pptx_coerce()
    except Exception as e:
        print(f"\n[FATAL] write_pptx coerce test crashed: {e}")
        traceback.print_exc()

    try:
        test_write_pdf_empty_table()
    except Exception as e:
        print(f"\n[FATAL] write_pdf empty table test crashed: {e}")
        traceback.print_exc()

    try:
        test_write_xlsx_data()
    except Exception as e:
        print(f"\n[FATAL] write_xlsx data test crashed: {e}")
        traceback.print_exc()

    try:
        test_write_pptx_layout()
    except Exception as e:
        print(f"\n[FATAL] write_pptx layout test crashed: {e}")
        traceback.print_exc()

    try:
        test_roundtrip_consistency()
    except Exception as e:
        print(f"\n[FATAL] Round-trip test crashed: {e}")
        traceback.print_exc()

    try:
        test_write_docx_exclusion()
    except Exception as e:
        print(f"\n[FATAL] write_docx exclusion test crashed: {e}")
        traceback.print_exc()

    try:
        test_read_docx_doc_cleanup()
    except Exception as e:
        print(f"\n[FATAL] read_docx doc cleanup test crashed: {e}")
        traceback.print_exc()

    # SUMMARY
    print("\n" + "=" * 70)
    print("DEEP DIVE SUMMARY")
    print("=" * 70)
    print(f"Total BUGS found:   {len(BUGS_FOUND)}")
    print(f"Total WARNINGS:     {len(WARNINGS)}")

    if BUGS_FOUND:
        print("\n--- ALL BUGS ---")
        for i, bug in enumerate(BUGS_FOUND, 1):
            print(f"\n  BUG #{i}: [{bug['test']}] {bug['description']}")
            if bug["details"]:
                for line in bug["details"].split("\n"):
                    print(f"           {line}")

    if WARNINGS:
        print("\n--- ALL WARNINGS ---")
        for i, w in enumerate(WARNINGS, 1):
            print(f"\n  WARN #{i}: [{w['test']}] {w['description']}")
            if w["details"]:
                for line in w["details"].split("\n"):
                    print(f"            {line}")

    print("\n" + "=" * 70)
    print("DONE")

# 编辑历史: 2026-07-18 小健 导入阶段设置的环境变量必须在模块加载后清理，避免污染后续测试
else:
    # 非 __main__ 运行（被 pytest 收集导入）时，清理临时 config 环境变量
    os.environ.pop("OMNIAGENT_CONFIG_PATH", None)
    print("=" * 70)
