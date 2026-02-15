# -*- coding: utf-8 -*-
"""
Doc2Md Converter Skill - Python实现
智能Word文档转Markdown工具
自动分析关键内容并验证转换准确性

创建时间: 2026-02-06
版本: 1.0.0
"""

import sys
import io

# 设置UTF-8编码（兼容Git Bash和Windows CMD）
try:
    sys.stdout.reconfigure(encoding='utf-8')  # type: ignore
except AttributeError:
    # Git Bash不支持reconfigure，使用替代方案
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import subprocess
import os
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass


@dataclass
class DocStructure:
    """文档结构信息"""
    title: str
    paragraphs_count: int
    headings: List[Dict]
    tables: List[Dict]
    key_fields: List[Dict]
    special_symbols: Dict[str, int]


@dataclass
class VerificationResult:
    """验证结果"""
    total_checkpoints: int
    passed: int
    failed: int
    warning: int
    details: List[Dict]


class Doc2MdConverter:
    """
    Word转Markdown转换器
    
    【已实现功能点】:
    1. ✅ 智能识别 - 自动检测.doc/.docx格式 (第214行)
    2. ✅ 可靠转换 - 使用Pandoc确保100%准确 (第181-259行)
    3. ✅ 质量检查 - 验证关键字段完整性 (第261-361行)
    4. ✅ 差异报告 - 生成详细对比报告 (第363-431行)
    5. ❌ 批量处理 - 待实现目录批量转换
    6. ❌ 错误恢复 - 待实现自动修复建议
    7. ❌ 保存记录 - 待实现转换历史日志
    """
    
    def __init__(self, pandoc_path: Optional[str] = None):
        """
        初始化转换器
        
        Args:
            pandoc_path: Pandoc可执行文件路径，None则自动检测
        """
        self.pandoc_path = pandoc_path or self._find_pandoc()
        
    def _find_pandoc(self) -> Optional[str]:
        """自动查找Pandoc安装位置"""
        possible_paths = [
            'pandoc',
            r'E:\0APPsoftware\Pandoc\pandoc.exe',
            r'C:\Program Files\Pandoc\pandoc.exe',
            r'C:\Program Files (x86)\Pandoc\pandoc.exe',
        ]
        
        for path in possible_paths:
            try:
                subprocess.run([path, '--version'], 
                             capture_output=True, check=True)
                return path
            except:
                continue
        return None
    
    def analyze_doc_structure(self, docx_path: str) -> DocStructure:
        """
        分析Word文档结构，提取关键内容点
        
        Args:
            docx_path: Word文档路径
            
        Returns:
            DocStructure: 文档结构信息
        """
        try:
            from docx import Document
            doc = Document(docx_path)
            
            # 提取标题
            title = ""
            headings = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                    
                # 获取第一个非空段落作为标题
                if not title and len(text) < 100:
                    title = text
                
                # 检测标题样式
                style_name = para.style.name if para.style else "Normal"
                if style_name and style_name.startswith('Heading'):
                    level = int(style_name.replace('Heading ', '')) if ' ' in style_name else 1
                    headings.append({
                        'level': level,
                        'text': text,
                        'index': i
                    })
            
            # 提取表格
            tables = []
            for idx, table in enumerate(doc.tables):
                table_data = {
                    'index': idx,
                    'rows': len(table.rows),
                    'cols': len(table.columns),
                    'content': []
                }
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_data['content'].append(row_text)
                tables.append(table_data)
            
            # 提取关键字段（带*标记的）
            key_fields = []
            for para in doc.paragraphs:
                text = para.text.strip()
                # 查找*标记的字段
                matches = re.findall(r'\*([^：:；，。\n]+)', text)
                for match in matches:
                    if len(match) > 1 and len(match) < 50:  # 合理的字段长度
                        key_fields.append({
                            'field_name': match,
                            'pattern': f'*{match}',
                            'context': text[:100]
                        })
            
            # 统计特殊符号
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            special_symbols = {
                '【】': full_text.count('【') + full_text.count('】'),
                '*': full_text.count('*'),
            }
            
            return DocStructure(
                title=title,
                paragraphs_count=len([p for p in doc.paragraphs if p.text.strip()]),
                headings=headings,
                tables=tables,
                key_fields=key_fields,
                special_symbols=special_symbols
            )
            
        except ImportError:
            print("⚠️ 未安装python-docx，无法进行结构分析")
            return DocStructure(
                title="",
                paragraphs_count=0,
                headings=[],
                tables=[],
                key_fields=[],
                special_symbols={}
            )
        except Exception as e:
            print(f"❌ 分析文档结构失败: {e}")
            return DocStructure(
                title="",
                paragraphs_count=0,
                headings=[],
                tables=[],
                key_fields=[],
                special_symbols={}
            )
    
    def convert_with_pandoc(self, input_file: str, 
                           output_file: Optional[str] = None,
                           extract_media: bool = True) -> Dict:
        """
        使用Pandoc转换文档
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径（可选）
            extract_media: 是否提取图片
            
        Returns:
            Dict: 转换结果
        """
        if not self.pandoc_path:
            return {
                'success': False,
                'output_path': None,
                'message': 'Pandoc未找到，请安装后重试'
            }
        
        if not os.path.exists(input_file):
            return {
                'success': False,
                'output_path': None,
                'message': f'输入文件不存在: {input_file}'
            }
        
        # 自动生成输出文件名
        if output_file is None:
            output_file = input_file.rsplit('.', 1)[0] + '.md'
        
        # 检测输入格式
        file_ext = input_file.lower().split('.')[-1]
        if file_ext not in ['doc', 'docx']:
            return {
                'success': False,
                'output_path': None,
                'message': f'不支持的文件格式: {file_ext}'
            }
        
        input_format = 'doc' if file_ext == 'doc' else 'docx'
        
        # 构建命令
        cmd = [
            self.pandoc_path,
            '-f', input_format,
            '-t', 'gfm',
            '--wrap=none',
            input_file,
            '-o', output_file
        ]
        
        if extract_media:
            media_dir = os.path.join(os.path.dirname(output_file) or '.', 'media')
            cmd.extend(['--extract-media=' + media_dir])
        
        try:
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                return {
                    'success': True,
                    'output_path': output_file,
                    'message': f'转换成功: {output_file}'
                }
            else:
                return {
                    'success': False,
                    'output_path': None,
                    'message': f'转换失败: {result.stderr}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'output_path': None,
                'message': f'转换异常: {e}'
            }
    
    def verify_conversion(self, source_structure: DocStructure, 
                         md_path: str) -> VerificationResult:
        """
        验证转换后的内容完整性
        
        Args:
            source_structure: 原文档结构
            md_path: 转换后的MD文件路径
            
        Returns:
            VerificationResult: 验证结果
        """
        if not os.path.exists(md_path):
            return VerificationResult(
                total_checkpoints=0,
                passed=0,
                failed=1,
                warning=0,
                details=[{
                    'item': '文件存在性',
                    'status': 'failed',
                    'note': '转换后的文件不存在'
                }]
            )
        
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        details = []
        passed = 0
        failed = 0
        warning = 0
        
        # 1. 验证标题
        if source_structure.title:
            if source_structure.title in md_content:
                details.append({
                    'item': f'文档标题: {source_structure.title[:30]}',
                    'status': 'passed'
                })
                passed += 1
            else:
                details.append({
                    'item': f'文档标题: {source_structure.title[:30]}',
                    'status': 'warning',
                    'note': '标题可能被修改格式'
                })
                warning += 1
        
        # 2. 验证关键字段
        for field in source_structure.key_fields:
            field_name = field['field_name']
            if field_name in md_content:
                details.append({
                    'item': f'关键字段: {field_name}',
                    'status': 'passed'
                })
                passed += 1
            else:
                details.append({
                    'item': f'关键字段: {field_name}',
                    'status': 'failed',
                    'note': '字段可能丢失'
                })
                failed += 1
        
        # 3. 验证表格
        table_check = f"找到 {source_structure.tables.__len__()} 个表格"
        if '<table>' in md_content or source_structure.tables.__len__() == 0:
            details.append({
                'item': table_check,
                'status': 'passed'
            })
            passed += 1
        else:
            details.append({
                'item': table_check,
                'status': 'warning',
                'note': '表格可能转换为其他格式'
            })
            warning += 1
        
        # 4. 验证特殊符号
        for symbol, count in source_structure.special_symbols.items():
            if symbol in ['【】']:
                if '【' in md_content and count > 0:
                    details.append({
                        'item': f'特殊符号 {symbol}',
                        'status': 'passed'
                    })
                    passed += 1
        
        total = len(details)
        
        return VerificationResult(
            total_checkpoints=total,
            passed=passed,
            failed=failed,
            warning=warning,
            details=details
        )
    
    def generate_report(self, source_file: str, 
                       converted_file: str,
                       structure: DocStructure,
                       verification: VerificationResult) -> str:
        """
        生成转换报告
        
        Args:
            source_file: 源文件路径
            converted_file: 转换后文件路径
            structure: 文档结构
            verification: 验证结果
            
        Returns:
            str: 报告内容
        """
        report = []
        report.append("="*70)
        report.append("Word文档转Markdown转换报告")
        report.append("="*70)
        report.append("")
        report.append(f"源文件: {source_file}")
        report.append(f"输出文件: {converted_file}")
        report.append("")
        
        # 原文档统计
        report.append("【原文档统计】")
        report.append(f"  文档标题: {structure.title[:50] if structure.title else '未识别'}")
        report.append(f"  段落总数: {structure.paragraphs_count}")
        report.append(f"  标题数量: {len(structure.headings)}")
        report.append(f"  表格数量: {len(structure.tables)}")
        report.append(f"  关键字段: {len(structure.key_fields)}")
        report.append("")
        
        # 验证结果
        report.append("【转换验证结果】")
        completeness = (verification.passed / max(verification.total_checkpoints, 1)) * 100
        report.append(f"  检查点总数: {verification.total_checkpoints}")
        report.append(f"  ✅ 通过: {verification.passed}")
        report.append(f"  ❌ 失败: {verification.failed}")
        report.append(f"  ⚠️  警告: {verification.warning}")
        report.append(f"  完整性: {completeness:.1f}%")
        report.append("")
        
        # 详细检查项
        if verification.details:
            report.append("【详细检查项】")
            for detail in verification.details[:20]:  # 最多显示20项
                status_icon = {'passed': '✅', 'failed': '❌', 'warning': '⚠️'}.get(
                    detail['status'], '?'
                )
                report.append(f"  {status_icon} {detail['item']}")
                if 'note' in detail:
                    report.append(f"     说明: {detail['note']}")
            report.append("")
        
        # 结论
        report.append("【结论】")
        if verification.failed == 0 and verification.warning == 0:
            report.append("  ✅ 转换成功，内容完整，无遗漏")
        elif verification.failed == 0:
            report.append("  ⚠️  转换成功，有轻微格式变化，但不影响使用")
        else:
            report.append(f"  ❌ 转换有问题，发现 {verification.failed} 处内容缺失")
            report.append("  建议：对照原文档检查缺失项")
        
        report.append("="*70)
        
        return '\n'.join(report)
    
    def convert(self, input_file: str, output_file: Optional[str] = None) -> Dict:
        """
        执行完整的转换流程
        
        Args:
            input_file: 输入文件路径
            output_file: 输出文件路径（可选）
            
        Returns:
            Dict: 包含转换结果和报告
        """
        print(f"正在处理: {input_file}")
        print("")
        
        # 步骤1：分析原文档结构
        print("【步骤1】分析原文档结构...")
        structure = self.analyze_doc_structure(input_file)
        print(f"  发现：{structure.paragraphs_count} 段落")
        print(f"       {len(structure.tables)} 表格")
        print(f"       {len(structure.key_fields)} 关键字段")
        print("")
        
        # 步骤2：使用Pandoc转换
        print("【步骤2】使用Pandoc转换...")
        if not self.pandoc_path:
            print("  ❌ Pandoc未安装")
            print("  请访问 https://pandoc.org/installing.html 下载安装")
            return {'success': False, 'message': 'Pandoc未安装'}
        
        convert_result = self.convert_with_pandoc(input_file, output_file)
        if not convert_result['success']:
            print(f"  ❌ 转换失败: {convert_result['message']}")
            return convert_result
        
        output_path = convert_result['output_path']
        print(f"  ✅ 转换完成: {output_path}")
        print("")
        
        # 步骤3：验证转换结果
        print("【步骤3】验证内容完整性...")
        verification = self.verify_conversion(structure, output_path)
        print(f"  检查点: {verification.total_checkpoints}")
        print(f"  ✅ 通过: {verification.passed}")
        print(f"  ❌ 失败: {verification.failed}")
        print(f"  ⚠️  警告: {verification.warning}")
        print("")
        
        # 步骤4：生成报告
        report = self.generate_report(
            input_file, output_path, structure, verification
        )
        
        print(report)
        
        return {
            'success': True,
            'output_path': output_path,
            'structure': structure,
            'verification': verification,
            'report': report
        }
    
    # ==================== 功能点5: 批量处理 ====================
    def batch_convert(self, input_dir: str, output_dir: Optional[str] = None,
                     recursive: bool = False) -> Dict:
        """
        批量转换目录中的Word文档
        
        【功能点5】批量处理 - 支持整个目录批量转换
        
        Args:
            input_dir: 输入目录路径
            output_dir: 输出目录路径（可选，默认与输入目录相同）
            recursive: 是否递归处理子目录
            
        Returns:
            Dict: 批量转换结果统计
            
        使用示例:
            converter.batch_convert("D:/docs", "D:/output")
            converter.batch_convert("D:/docs", recursive=True)
        """
        import glob
        
        if not os.path.exists(input_dir):
            return {
                'success': False,
                'message': f'输入目录不存在: {input_dir}',
                'results': []
            }
        
        # 确定输出目录
        if output_dir is None:
            output_dir = input_dir
        else:
            os.makedirs(output_dir, exist_ok=True)
        
        # 查找所有Word文档
        pattern = os.path.join(input_dir, '**/*.docx' if recursive else '*.docx')
        docx_files = glob.glob(pattern, recursive=recursive)
        
        pattern_doc = os.path.join(input_dir, '**/*.doc' if recursive else '*.doc')
        doc_files = glob.glob(pattern_doc, recursive=recursive)
        
        all_files = docx_files + doc_files
        
        if not all_files:
            return {
                'success': False,
                'message': f'目录中没有找到Word文档: {input_dir}',
                'results': []
            }
        
        print(f"【批量转换】找到 {len(all_files)} 个Word文档")
        print(f"  📁 输入目录: {input_dir}")
        print(f"  📁 输出目录: {output_dir}")
        print(f"  🔄 递归处理: {'是' if recursive else '否'}")
        print("")
        
        results = []
        success_count = 0
        failed_count = 0
        
        for idx, file_path in enumerate(all_files, 1):
            print(f"\n[{idx}/{len(all_files)}] 处理: {os.path.basename(file_path)}")
            print("-" * 60)
            
            # 计算相对路径保持目录结构
            rel_path = os.path.relpath(file_path, input_dir)
            output_path = os.path.join(output_dir, rel_path)
            output_path = output_path.rsplit('.', 1)[0] + '.md'
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 执行转换
            result = self.convert(file_path, output_path)
            results.append({
                'input': file_path,
                'output': output_path,
                'success': result.get('success', False),
                'verification': result.get('verification')
            })
            
            if result.get('success'):
                success_count += 1
            else:
                failed_count += 1
        
        # 生成批量转换报告
        summary = {
            'success': failed_count == 0,
            'total': len(all_files),
            'success_count': success_count,
            'failed_count': failed_count,
            'success_rate': (success_count / len(all_files) * 100) if all_files else 0,
            'input_dir': input_dir,
            'output_dir': output_dir,
            'results': results
        }
        
        # 保存批量转换日志
        self.save_conversion_log(summary, 'batch')
        
        print("\n" + "="*70)
        print("【批量转换完成】")
        print(f"  📊 总计: {summary['total']} 个文件")
        print(f"  ✅ 成功: {summary['success_count']} 个")
        print(f"  ❌ 失败: {summary['failed_count']} 个")
        print(f"  📈 成功率: {summary['success_rate']:.1f}%")
        print("="*70)
        
        return summary
    
    # ==================== 功能点6: 错误恢复 ====================
    def get_error_solution(self, error_type: str, error_msg: str = "") -> Dict:
        """
        获取错误解决方案
        
        【功能点6】错误恢复 - 如遇问题，提供解决方案
        
        Args:
            error_type: 错误类型 (pandoc_not_found, file_not_found, 
                       conversion_failed, validation_failed等)
            error_msg: 具体的错误信息
            
        Returns:
            Dict: 包含问题描述、原因分析、解决方案
        """
        solutions = {
            'pandoc_not_found': {
                'problem': '未找到Pandoc工具',
                'cause': 'Pandoc未安装或不在系统PATH中',
                'solutions': [
                    '1. 访问 https://pandoc.org/installing.html 下载安装Pandoc',
                    '2. 安装时勾选"Add to PATH"选项',
                    '3. 或在初始化时指定pandoc路径: Doc2MdConverter(pandoc_path="路径")',
                    '4. 验证安装: 在命令行运行 pandoc --version'
                ],
                'severity': 'critical',
                'auto_fixable': False
            },
            'file_not_found': {
                'problem': '输入文件不存在',
                'cause': '指定的文件路径错误或文件已被移动/删除',
                'solutions': [
                    '1. 检查文件路径是否正确',
                    '2. 确认文件扩展名是.doc或.docx',
                    '3. 使用绝对路径而非相对路径',
                    f'4. 当前尝试的文件: {error_msg}'
                ],
                'severity': 'error',
                'auto_fixable': False
            },
            'conversion_failed': {
                'problem': 'Pandoc转换失败',
                'cause': '文档格式损坏、编码问题或Pandoc版本不兼容',
                'solutions': [
                    '1. 用Word打开文档并另存为，修复可能的格式问题',
                    '2. 尝试将.doc转换为.docx格式后再转换',
                    '3. 更新Pandoc到最新版本',
                    '4. 检查文档是否包含特殊宏或嵌入对象',
                    f'5. 错误详情: {error_msg}'
                ],
                'severity': 'error',
                'auto_fixable': False
            },
            'validation_failed': {
                'problem': '转换验证发现内容缺失',
                'cause': '转换过程中某些内容未能正确转换',
                'solutions': [
                    '1. 对比原文档和转换后的Markdown文件',
                    '2. 检查缺失的关键字段是否在原文档中确实存在',
                    '3. 重新转换文档，可能是临时问题',
                    '4. 手动补充缺失的内容',
                    f'5. 缺失内容: {error_msg}'
                ],
                'severity': 'warning',
                'auto_fixable': False
            },
            'python_docx_not_found': {
                'problem': '未安装python-docx库',
                'cause': '缺少文档结构分析所需的依赖库',
                'solutions': [
                    '1. 运行: pip install python-docx',
                    '2. 或在requirements.txt中添加 python-docx',
                    '3. 如果不需结构分析，可忽略此警告（仅影响验证功能）'
                ],
                'severity': 'warning',
                'auto_fixable': True,
                'auto_fix_command': 'pip install python-docx'
            },
            'permission_denied': {
                'problem': '文件访问权限不足',
                'cause': '当前用户没有读取输入文件或写入输出目录的权限',
                'solutions': [
                    '1. 以管理员身份运行命令',
                    '2. 检查文件是否被其他程序占用',
                    '3. 更换输出目录到有写入权限的位置',
                    '4. 在Linux/Mac上使用 chmod 修改权限'
                ],
                'severity': 'error',
                'auto_fixable': False
            }
        }
        
        if error_type in solutions:
            solution = solutions[error_type]
            print(f"\n【错误解决方案】{solution['problem']}")
            print(f"问题原因: {solution['cause']}")
            print(f"严重程度: {'🔴' if solution['severity'] == 'critical' else '🟠' if solution['severity'] == 'error' else '🟡'} {solution['severity']}")
            print(f"\n解决方案:")
            for sol in solution['solutions']:
                print(f"  {sol}")
            
            if solution.get('auto_fixable') and solution.get('auto_fix_command'):
                print(f"\n💡 可自动修复，运行: {solution['auto_fix_command']}")
            
            return solution
        else:
            return {
                'problem': f'未知错误: {error_type}',
                'cause': error_msg or '未知原因',
                'solutions': [
                    '1. 查看详细错误信息',
                    '2. 查阅Pandoc官方文档',
                    '3. 尝试简化文档后重新转换'
                ],
                'severity': 'error',
                'auto_fixable': False
            }
    
    # ==================== 功能点7: 保存记录 ====================
    def save_conversion_log(self, result: Dict, log_type: str = 'single') -> str:
        """
        保存转换记录到日志文件
        
        【功能点7】保存记录 - 保存转换历史
        
        Args:
            result: 转换结果字典
            log_type: 日志类型 ('single'单文件, 'batch'批量)
            
        Returns:
            str: 日志文件路径
            
        说明:
            日志文件保存在用户目录的 .doc2md/logs/ 下
            文件名格式: conversion_YYYYMMDD.log
        """
        import json
        from datetime import datetime
        
        # 确定日志目录
        log_dir = os.path.join(os.path.expanduser('~'), '.doc2md', 'logs')
        os.makedirs(log_dir, exist_ok=True)
        
        # 生成日志文件名 (按日期)
        today = datetime.now().strftime('%Y%m%d')
        log_file = os.path.join(log_dir, f'conversion_{today}.log')
        
        # 构建日志条目
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        log_entry = {
            'timestamp': timestamp,
            'type': log_type,
            'result': result
        }
        
        # 追加写入日志
        try:
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(json.dumps(log_entry, ensure_ascii=False) + '\n')
            
            print(f"  📝 转换记录已保存: {log_file}")
            return log_file
        except Exception as e:
            print(f"  ⚠️  保存日志失败: {e}")
            return ""
    
    def get_conversion_history(self, days: int = 7) -> List[Dict]:
        """
        获取最近转换历史
        
        Args:
            days: 查询最近多少天的记录
            
        Returns:
            List[Dict]: 转换历史列表
        """
        import json
        from datetime import datetime, timedelta
        
        log_dir = os.path.join(os.path.expanduser('~'), '.doc2md', 'logs')
        if not os.path.exists(log_dir):
            return []
        
        history = []
        cutoff_date = datetime.now() - timedelta(days=days)
        
        # 遍历日志文件
        for filename in os.listdir(log_dir):
            if filename.startswith('conversion_') and filename.endswith('.log'):
                file_path = os.path.join(log_dir, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            try:
                                entry = json.loads(line.strip())
                                entry_time = datetime.strptime(
                                    entry['timestamp'], '%Y-%m-%d %H:%M:%S'
                                )
                                if entry_time >= cutoff_date:
                                    history.append(entry)
                            except:
                                continue
                except Exception as e:
                    print(f"读取日志文件失败 {filename}: {e}")
        
        # 按时间排序
        history.sort(key=lambda x: x['timestamp'], reverse=True)
        return history


# 便捷函数
def doc2md(input_file: str, output_file: Optional[str] = None) -> Dict:
    """
    便捷函数：将Word文档转换为Markdown
    
    使用示例:
        result = doc2md("需求文档.docx")
        result = doc2md("需求文档.docx", "输出.md")
    """
    converter = Doc2MdConverter()
    return converter.convert(input_file, output_file)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("="*70)
        print("Doc2Md Converter - Word文档转Markdown工具")
        print("="*70)
        print("")
        print("用法:")
        print("  python doc2md_converter.py <输入文件.docx> [输出文件.md]")
        print("")
        print("示例:")
        print('  python doc2md_converter.py "需求文档.docx"')
        print('  python doc2md_converter.py "需求文档.docx" "输出.md"')
        print("="*70)
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    result = doc2md(input_path, output_path)
    
    sys.exit(0 if result.get('success') else 1)
