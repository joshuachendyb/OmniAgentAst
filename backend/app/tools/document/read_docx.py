# -*- coding: utf-8 -*-
# 编辑历史:
# 2026-07-20 - 小欧 - 自然单位翻页 feat:
#   1. read_docx 增加 offset/limit/tail 参数
#   2. 新增行翻页参数校验逻辑
#   3. 复用 line_pager.select_lines 行窗口
#   4. llm_data 增加 total_lines 指标
#   5. 支持 warning 状态(exec_code/detail/hint)
# 2026-07-21 - 小欧 - 入参即信任: _build_read_docx_llm_data 加 user_limit 参数, 入 action.params, 支撑 formatter 动态调行数上限
# 2026-07-21 - 小欧 - 文件大小安全检测: 入口加 READ_DOCX_INPUT_MAX_BYTES 字节检查，超限拒读防OOM
# 2026-07-23 - 小欧 - 三堂会审5bug修复: outlimit len(text)截断后求值+删死import os as _os_mod
# 2026-07-24 - 小欧 - 修复: error summary嵌入full detail → 改用truncate_summary(detail)首行
# 2026-07-26 - 小欧 - OOD: 确认READ_DOCX_INPUT_MAX_BYTES未落地,OOM自然抛出被except捕获(同dataanalysis模式); 删doc_path多余变量(KISS-DIRECT)
# 2026-07-26 - 小欧 - 清理: 删logger死import(全文件无logger调用)
# 2026-08-13 - 小欧 - A5职责拆分: hint_* 错误提示函数/导入源改 app.tools.toolhelper.error_hints
"""
D2: read_docx — 读取Word文档

从document_tools.py拆分而来 — 小欧 2026-06-22
"""
# 【铁规1】helper/被调函数(以下划线_开头的函数)只返回raw dict，严禁调用build_success/build_error/build_warning和构建llm_data。
# build3+llm_data只能在tool的main函数(对外公开的函数)中包装。违反此规则的代码视为不合规。
# 【铁规2】工具返回原始data，禁止调用truncate_data_for_frontend。截断只能在前端yield层。
# 【铁规3】计时(duration_ms计算)只能在tool的主函数中，严禁在子函数/helper中计时。

import time as _time_mod
from pathlib import Path
from typing import Any, Dict, Optional

from app.tools.tool_response import build_success, build_error, build_warning
from app.tools.tool_fc_helper import _check_module
from app.tools.validate.file_type_checker import check_for_document_tool
from app.tools.toolhelper.error_hints import hint_for_read_error
from app.tools.tool_constants import ERR_DOC_READ_DOCX, READ_DOCX_OUTLIMIT_CHARS
from app.tools.toolhelper.line_pager import select_lines
from app.utils.text_utils import truncate_summary


def _build_read_docx_llm_data(
    exec_code: str, duration_ms: int,
    file_path: str = "", para_count: int = 0, text_len: int = 0,
    non_empty: int = 0, empty: int = 0, table_count: int = 0,
    total_lines: int = 0, detail: str = "", hint: str = "",
    user_limit: Optional[int] = None,
) -> Dict[str, Any]:
    """read_docx的llm_data构建函数 — 小健 2026-06-21 — 小欧 2026-06-22 — 小欧 2026-07-05 加hint参数 — 小欧 2026-07-06 丰富summary
    2026-07-21 入参即信任: 补 user_limit 写入 action.params — 小欧"""
    _act_params = {"file_path": file_path}
    if user_limit is not None:
        _act_params["limit"] = user_limit
    if exec_code == "error":
        _err_summary = truncate_summary(detail)
        return {
            "summary": f"读取Word{file_path}，失败" + (f": {_err_summary}" if _err_summary else ""),
            "action": {"tool": "read_docx", "tool_zh": "读取Word", "target": file_path, "params": _act_params},
            "status": {"exec_code": "error", "message": "读取Word失败", "code": ERR_DOC_READ_DOCX, "detail": detail, "hint": hint if hint else "读取失败,详见错误明细"},
            "duration_ms": duration_ms,
            "metrics": {},
        }
    # summary: 段落数(含非空/空分段)、字符数、表格数 — 小欧 2026-07-06
    parts = []
    if empty > 0:
        parts.append(f"{para_count}段({non_empty}非空, {empty}空)")
    else:
        parts.append(f"{para_count}段")
    parts.append(f"{text_len}字符")
    if table_count:
        parts.append(f"{table_count}项表格")
    summary_str = f"读取Word{file_path}，成功: " + "，".join(parts)
    return {
        "summary": summary_str,
        "action": {"tool": "read_docx", "tool_zh": "读取Word", "target": file_path, "params": _act_params},
        "status": {"exec_code": exec_code, "message": "读取Word成功" if exec_code == "success" else "读取Word有警告", "code": "", "detail": detail, "hint": hint},
        "duration_ms": duration_ms,
        "metrics": {
            "para_count": {"value": para_count, "text": f"{para_count}段"},
            "text_len": {"value": text_len, "text": f"{text_len}字符"},
            "total_lines": {"value": total_lines, "text": f"{total_lines}行"},
        },
    }


def read_docx(
    path: str,
    offset: Optional[int] = None,
    limit: Optional[int] = None,
    tail: Optional[int] = None,
) -> Dict[str, Any]:
    """读取Word(.docx)文档 — 小沈 2026-06-19 — 小欧 2026-06-22 独立文件 — 小欧 2026-06-24 增加文件类型前置检查 — 小欧 2026-06-24 移除.doc死代码(pandoc转换)"""
    t0 = _time_mod.perf_counter()
    file_path = path

    # 文件类型前置检查 — 小欧 2026-06-24
    is_valid, error_detail, suggested_tool = check_for_document_tool(path)
    if not is_valid:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        if suggested_tool:
            _hint = f"建议使用{suggested_tool}工具"
        elif suggested_tool == "":
            _hint = "请检查文件路径和文件名是否正确"
        else:
            _hint = "文件类型不匹配,请使用正确的文档格式"
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=error_detail, hint=_hint)
        return build_error(data={}, llm_data=llm_data)

    if not _check_module("docx"):
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail="python-docx库未安装", hint="请安装python-docx库")
        return build_error(data={}, llm_data=llm_data)

    # 行翻页参数校验(自然单位治理 2026-07-20 小欧, 复用 readtext 的 offset/limit/tail 语义)
    if limit is not None and (limit < 1 or limit > 1000):
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=f"limit参数必须在1-1000之间,传入值: {limit}", hint="limit参数必须设置在1-1000之间")
        return build_error(data={}, llm_data=llm_data)
    if tail is not None and tail < 1:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=f"tail参数不能小于1,传入值: {tail}", hint="tail参数不能小于1")
        return build_error(data={}, llm_data=llm_data)
    if tail is not None and (offset is not None or limit is not None):
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail="tail参数不能与offset/limit同时使用", hint="tail与offset/limit参数互斥,请选择其一")
        return build_error(data={}, llm_data=llm_data)
    if offset is not None:
        if offset < 1:
            duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
            llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=f"offset参数必须>=1,当前值: {offset}", hint="offset行号从1开始")
            return build_error(data={}, llm_data=llm_data)
        if limit is None:
            duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
            llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=f"offset参数必须同时提供limit参数,当前offset={offset}", hint="请提供limit参数配合offset")
            return build_error(data={}, llm_data=llm_data)

    try:
        import docx

        doc = docx.Document(Path(file_path))
        paragraphs = [para.text for para in doc.paragraphs]
        non_empty_paragraphs = [p for p in paragraphs if p.strip()]
        text = "\n".join(non_empty_paragraphs)
        # outlimit: 仅全量读取(无翻页参数)截断, 翻页由用户参数控制
        if text and offset is None and limit is None and tail is None:
            _orig_len = len(text)
            if _orig_len > READ_DOCX_OUTLIMIT_CHARS:
                text = text[:READ_DOCX_OUTLIMIT_CHARS] + \
                    f"\n... (内容已截断: 原文{_orig_len}字符, 保留{READ_DOCX_OUTLIMIT_CHARS}字符) ...\n"
        empty_para_count = len(paragraphs) - len(non_empty_paragraphs)

        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append(table_rows)

        # —— 自然单位治理(2026-07-20 小欧): Word 无可靠页码, 按段落/行翻页(offset/limit/tail), 复用 line_pager.select_lines ——
        lines = text.split("\n")
        _sel = select_lines(lines, offset, limit, tail)
        selected_text = _sel["content"]
        total_lines = _sel["total_lines"]
        warning = _sel.get("warning")

        result_data = {"text": selected_text}
        if tables_data:
            result_data["tables"] = tables_data

        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        if warning:
            llm_data = _build_read_docx_llm_data(
                "warning", duration_ms, path, len(paragraphs), len(selected_text),
                len(non_empty_paragraphs), empty_para_count, len(tables_data),
                total_lines, detail=warning, hint="请调整offset/limit参数",
                user_limit=limit,
            )
            return build_warning(data=result_data, llm_data=llm_data)
        llm_data = _build_read_docx_llm_data(
            "success", duration_ms, path, len(paragraphs), len(selected_text),
            len(non_empty_paragraphs), empty_para_count, len(tables_data), total_lines,
            user_limit=limit,
        )
        # =============================================================================
        # 数据设计：paragraph_count/non_empty_paragraph_count/empty_paragraph_count/
        # table_count/total_lines 从 data 移除，通过 llm_data.metrics + summary 传递给 LLM
        # summary 示例: "读取Word成功: 5段(3非空, 2空), 1200字符, 1项表格"
        # data 只保留 text/tables 纯数据 (formatter 渲染用)
        # — 小欧 2026-07-06; 2026-07-20 加 offset/limit/tail 翻页(total_lines 入 metrics)
        # =============================================================================
        # ---- observation_formatter route -------------------------------------------
        # branch: #10b raw prose (DOCX/文本)
        # trigger: "text" in data and isinstance(data["text"], str)
        # handler: _format_prose_result(data) — 两态说明行(已截断/完整) + 正文 + 额外字段
        # file:    observation_formatter.py:_format_prose_result
        # ------------------------------------------------------------------------------
        return build_success(data=result_data, llm_data=llm_data)
    except Exception as e:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_read_docx_llm_data("error", duration_ms, path, detail=str(e), hint=hint_for_read_error(e, path))
        return build_error(data={}, llm_data=llm_data)
