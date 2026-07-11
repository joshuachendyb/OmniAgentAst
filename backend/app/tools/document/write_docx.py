# -*- coding: utf-8 -*-
"""
D5: write_docx — 写入Word文档

从document_tools.py拆分而来 — 小欧 2026-06-22

"""
# 【铁规1】helper/被调函数(以下划线_开头的函数)只返回raw dict，严禁调用build_success/build_error/build_warning和构建llm_data。
# build3+llm_data只能在tool的main函数(对外公开的函数)中包装。违反此规则的代码视为不合规。
# 【铁规2】工具返回原始data，禁止调用truncate_data_for_frontend。截断只能在前端yield层。
# 【铁规3】计时(duration_ms计算)只能在tool的主函数中，严禁在子函数/helper中计时。

import re
import time as _time_mod
from pathlib import Path
from typing import Any, Dict, List, Optional

from app.tools.tool_response import build_success, build_error
from app.tools.tool_fc_helper import _check_module
from app.tools.validate.file_type_checker import check_office_file
from app.tools.validate.file_safety_checker import check_content_safety
from app.tools.tool_constants import ERR_WRITE_DOCX
from app.tools.validate.file_path_checker import permission_error_hint, hint_for_write_error
from app.logger import logger
from app.utils.table_helper import parse_markdown_table, calculate_column_widths, get_table_header_style_config, normalize_table_data
from app.tools.document.md_inline_utils import _parse_inline_md


def _apply_inline_formatting(p, text):
    """给python-docx段落添加行内格式run — 小欧 2026-07-08"""
    from docx.shared import Pt
    for seg_text, bold, italic, code, link_url in _parse_inline_md(text):
        run = p.add_run(seg_text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)


def _format_cell_content(cell, text):
    """给表格单元格添加行内格式文本 — 小欧 2026-07-08"""
    cell.text = ""
    _apply_inline_formatting(cell.paragraphs[0], str(text))


def _set_docx_table_style(table):
    """设置表格边框和表头样式 — 小健 2026-06-24"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor, Pt
    
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    header_config = get_table_header_style_config()
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = header_config["bold"]
                    run.font.size = Pt(header_config["font_size"])
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), header_config["bg_color"])
            tcPr.append(shd)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)


def _set_docx_column_widths(table, table_data):
    """设置列宽自适应 — 小健 2026-06-24"""
    from docx.shared import Inches
    
    if not table_data or not table_data[0]:
        return
    
    col_widths = calculate_column_widths(table_data, total_width=6.0)
    for ci, width in enumerate(col_widths):
        if ci < len(table.columns):
            for cell in table.columns[ci].cells:
                cell.width = Inches(width)


def _build_write_docx_llm_data(
    exec_code: str, duration_ms: int,
    file_path: str = "", detail: str = "", user_title: str = "", hint: str = "",
) -> Dict[str, Any]:
    """write_docx的llm_data构建函数 — 小欧 2026-06-22 — 小欧 2026-07-05 加hint参数"""
    _act_params = {"file_path": file_path}
    if user_title:
        _act_params["title"] = user_title
    if exec_code == "error":
        return {
            "summary": f"写入Word{file_path}，失败: {detail}",
            "action": {"tool": "write_docx", "tool_zh": "写入Word", "target": file_path, "params": _act_params},
            "status": {"exec_code": "error", "message": "写入Word失败", "code": ERR_WRITE_DOCX, "detail": detail, "hint": hint if hint else "请检查路径和权限"},
            "duration_ms": duration_ms,
            "metrics": {},
        }
    return {
        "summary": f"写入Word{file_path}，成功",
        "action": {"tool": "write_docx", "tool_zh": "写入Word", "target": file_path, "params": _act_params},
        "status": {"exec_code": "success", "message": "写入Word成功", "code": "", "detail": "", "hint": ""},
        "duration_ms": duration_ms,
        "metrics": {},
    }




def write_docx(
    path: str,
    title: Optional[str] = None,
    content: Optional[str] = None,
    table_data: Optional[List[List[str]]] = None,
) -> Dict[str, Any]:
    """写入Word文档 — 小健 2026-06-24 支持Markdown表格+table_data互斥 — 小欧 2026-06-24 增加文件类型前置检查"""
    t0 = _time_mod.perf_counter()

    # 文件类型前置检查（含路径检查+类型检查+模块安全检查）— 北京老陈 2026-07-09
    is_valid, error_detail, hint = check_office_file(path, allow_create=True)
    if not is_valid:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_write_docx_llm_data("error", duration_ms, path, detail=error_detail, user_title=title or "", hint=hint)
        return build_error(data={}, llm_data=llm_data)

    if content is not None:
        cs_error, _ = check_content_safety(content, "docx")
        if cs_error:
            duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
            llm_data = _build_write_docx_llm_data("error", duration_ms, path, detail=cs_error, user_title=title or "", hint="请检查content参数")
            return build_error(data={}, llm_data=llm_data)

    if not _check_module("docx"):
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_write_docx_llm_data("error", duration_ms, path, detail="python-docx库未安装", user_title=title or "", hint="请安装python-docx库")
        return build_error(data={}, llm_data=llm_data)

    try:
        from docx import Document

        doc = Document()

        if title:
            doc.add_heading(title, 0)

        if content:
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].rstrip()
                if not line:
                    i += 1
                    continue
                
                p = None
                if line.startswith('# '):
                    p = doc.add_paragraph(style='Heading 1')
                    _apply_inline_formatting(p, line[2:])
                elif line.startswith('## '):
                    p = doc.add_paragraph(style='Heading 2')
                    _apply_inline_formatting(p, line[3:])
                elif line.startswith('### '):
                    p = doc.add_paragraph(style='Heading 3')
                    _apply_inline_formatting(p, line[4:])
                elif line.startswith('#### '):
                    p = doc.add_paragraph(style='Heading 4')
                    _apply_inline_formatting(p, line[5:])
                elif line.startswith('##### '):
                    p = doc.add_paragraph(style='Heading 5')
                    _apply_inline_formatting(p, line[6:])
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    _apply_inline_formatting(p, line[2:])
                elif re.match(r'^\d+\.\s', line):
                    p = doc.add_paragraph(style='List Number')
                    _apply_inline_formatting(p, re.sub(r'^\d+\.\s', '', line))
                elif line.startswith('|') and '|' in line[1:]:
                    table_rows, i = parse_markdown_table(lines, i)
                    if table_rows and table_rows[0]:
                        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        for ri, row_data in enumerate(table_rows):
                            for ci, cell_text in enumerate(row_data):
                                _format_cell_content(t.rows[ri].cells[ci], cell_text)
                        _set_docx_table_style(t)
                        _set_docx_column_widths(t, table_rows)
                    continue
                else:
                    p = doc.add_paragraph()
                    _apply_inline_formatting(p, line)
                i += 1
        
        if table_data:
            table_data = normalize_table_data(table_data)
            if table_data and len(table_data) > 0 and len(table_data[0]) > 0:
                t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for ri, row_data in enumerate(table_data):
                    for ci, cell_text in enumerate(row_data):
                        _format_cell_content(t.rows[ri].cells[ci], cell_text)
                _set_docx_table_style(t)
                _set_docx_column_widths(t, table_data)

        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)

        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        llm_data = _build_write_docx_llm_data("success", duration_ms, str(path), user_title=title or "")
        # =============================================================================
        # 数据设计：file_path 从 data 移除，通过 llm_data.summary 传入 LLM observation。
        # summary 已包含文件路径: "写入Word成功: /path.docx"
        # data 为空 dict 时 formatter 不追加详情，避免冗余。
        # — 小欧 2026-07-06
        # =============================================================================
        return build_success(data={}, llm_data=llm_data)
    except PermissionError as e:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        hint = permission_error_hint(path)
        llm_data = _build_write_docx_llm_data("error", duration_ms, path, detail=str(e), user_title=title or "", hint=hint)
        return build_error(data={}, llm_data=llm_data)
    except Exception as e:
        duration_ms = int((_time_mod.perf_counter() - t0) * 1000)
        hint = hint_for_write_error(e, path, "写入Word异常,请检查磁盘空间和权限")
        llm_data = _build_write_docx_llm_data("error", duration_ms, path, detail=str(e), user_title=title or "", hint=hint)
        return build_error(data={}, llm_data=llm_data)