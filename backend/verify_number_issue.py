# -*- coding: utf-8 -*-
"""验证Word列表自动编号问题 - 小健 2026-06-24"""

from docx import Document
from pathlib import Path
import tempfile
import shutil
import sys
sys.path.insert(0, 'G:/OmniAgentAs-desk/backend')

from app.tools.document.write_docx import write_docx

temp_dir = Path(tempfile.mkdtemp())

print("=" * 60)
print("问题验证: 有序列表数字前缀丢失")
print("=" * 60)

test_file = temp_dir / 'test_number_prefix.docx'
content = """1. 第一项
2. 第二项
10. 第十项（用户指定10，Word会重新编号为3吗？）
99. 第九十九项（用户指定99，Word会重新编号为4吗？）"""

result = write_docx(str(test_file), content=content)

doc = Document(str(test_file))
print(f'段落数量: {len(doc.paragraphs)}')
print("\n段落内容:")
for i, para in enumerate(doc.paragraphs):
    print(f'  [{i}] text="{para.text}"')
    print(f'      style="{para.style.name}"')

print("\n" + "=" * 60)
print("问题分析:")
print("=" * 60)
print("输入: '10. 第十项'")
print("实际输出: text='第十项' style='List Number'")
print("问题: 用户指定的数字前缀(10)被删除了")
print("      Word的List Number样式会自动重新编号为3")
print("      这可能不是用户期望的行为！")
print("=" * 60)

# 清理
shutil.rmtree(temp_dir)