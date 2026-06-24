# -*- coding: utf-8 -*-
"""验证write_docx实际内容 - 小健 2026-06-24"""

from docx import Document
from pathlib import Path
import tempfile
import shutil
import sys
sys.path.insert(0, 'G:/OmniAgentAs-desk/backend')

from app.tools.document.write_docx import write_docx

temp_dir = Path(tempfile.mkdtemp())

print("=" * 60)
print("测试1: 有序列表 - 验证内容和样式")
print("=" * 60)

test_file = temp_dir / 'test_ordered_list.docx'
content = """1. 第一项
2. 第二项
10. 第十项
99. 第九十九项"""

result = write_docx(str(test_file), content=content)

doc = Document(str(test_file))
print(f'段落数量: {len(doc.paragraphs)}')
for i, para in enumerate(doc.paragraphs):
    print(f'段落{i}: text="{para.text}" style="{para.style.name}"')

print("\n" + "=" * 60)
print("测试2: 无序列表 - 验证内容和样式")
print("=" * 60)

test_file2 = temp_dir / 'test_unordered_list.docx'
content2 = """- 项目1
- 项目2
* 星号项1
* 星号项2"""

result2 = write_docx(str(test_file2), content=content2)

doc2 = Document(str(test_file2))
print(f'段落数量: {len(doc2.paragraphs)}')
for i, para in enumerate(doc2.paragraphs):
    print(f'段落{i}: text="{para.text}" style="{para.style.name}"')

print("\n" + "=" * 60)
print("测试3: 标题 - 验证级别")
print("=" * 60)

test_file3 = temp_dir / 'test_headings.docx'
content3 = """# 一级标题
## 二级标题
### 三级标题
#### 四级标题
##### 五级标题"""

result3 = write_docx(str(test_file3), content=content3)

doc3 = Document(str(test_file3))
print(f'段落数量: {len(doc3.paragraphs)}')
for i, para in enumerate(doc3.paragraphs):
    print(f'段落{i}: text="{para.text}" style="{para.style.name}"')

# 清理
shutil.rmtree(temp_dir)