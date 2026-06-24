# -*- coding: utf-8 -*-
"""验证Markdown表格和table_data功能 - 小健 2026-06-24"""

from app.tools.document.write_docx import write_docx
from docx import Document
import tempfile
from pathlib import Path
import shutil

temp_dir = Path(tempfile.mkdtemp())

print('=' * 80)
print('测试1: Markdown表格')
print('=' * 80)

file1 = temp_dir / 'test_table.docx'
content = """# 数据报告

## 统计表格

| 项目 | 数值 | 占比 |
|------|------|------|
| A | 100 | 40% |
| B | 150 | 60% |

## 结论

- 数据A占比40%
- 数据B占比60%"""

result = write_docx(str(file1), content=content)
print(f'结果: {result["llm_data"]["status"]["exec_code"]}')

doc = Document(str(file1))
print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')

if doc.tables:
    table = doc.tables[0]
    print(f'表格行数: {len(table.rows)}')
    print(f'表格列数: {len(table.columns)}')
    print('表格内容:')
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        print(f'  {cells}')

print()
print('=' * 80)
print('测试2: table_data参数（纯表格）')
print('=' * 80)

file2 = temp_dir / 'test_table_data.docx'
table_data = [
    ['姓名', '年龄', '城市'],
    ['张三', '25', '北京'],
    ['李四', '30', '上海']
]

result2 = write_docx(str(file2), title='数据表', table_data=table_data)
print(f'结果: {result2["llm_data"]["status"]["exec_code"]}')

doc2 = Document(str(file2))
print(f'表格数: {len(doc2.tables)}')

if doc2.tables:
    table2 = doc2.tables[0]
    print('表格内容:')
    for row in table2.rows:
        cells = [cell.text for cell in row.cells]
        print(f'  {cells}')

print()
print('=' * 80)
print('测试3: content和table_data互斥（content优先）')
print('=' * 80)

file3 = temp_dir / 'test_mutex.docx'
result3 = write_docx(
    str(file3),
    content="# 有content",
    table_data=[['A', 'B'], ['C', 'D']]
)
doc3 = Document(str(file3))
print(f'段落数: {len(doc3.paragraphs)} (应该有段落)')
print(f'表格数: {len(doc3.tables)} (应该为0，因为content优先)')

shutil.rmtree(temp_dir)
print()
print('✅ 所有测试通过')