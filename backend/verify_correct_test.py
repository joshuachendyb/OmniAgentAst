# -*- coding: utf-8 -*-
"""正确的测试 - 验证实际内容 - 小健 2026-06-24"""

from docx import Document
from pathlib import Path
import tempfile
import shutil
import sys
sys.path.insert(0, 'G:/OmniAgentAs-desk/backend')

from app.tools.document.write_docx import write_docx

temp_dir = Path(tempfile.mkdtemp())

print("=" * 60)
print("正确测试1: 验证标题级别")
print("=" * 60)

test_file = temp_dir / 'test_heading_levels.docx'
content = """# 一级标题
## 二级标题
### 三级标题"""

write_docx(str(test_file), content=content)
doc = Document(str(test_file))

# 验证标题级别
assert len(doc.paragraphs) == 3, "应该有3个段落"
assert doc.paragraphs[0].style.name == 'Heading 1', "第一个应该是Heading 1"
assert doc.paragraphs[1].style.name == 'Heading 2', "第二个应该是Heading 2"
assert doc.paragraphs[2].style.name == 'Heading 3', "第三个应该是Heading 3"
print("✅ 标题级别验证通过")

print("\n" + "=" * 60)
print("正确测试2: 验证无序列表样式")
print("=" * 60)

test_file2 = temp_dir / 'test_bullet_list.docx'
content2 = """- 项目1
- 项目2
* 项目3"""

write_docx(str(test_file2), content=content2)
doc2 = Document(str(test_file2))

assert len(doc2.paragraphs) == 3, "应该有3个段落"
for para in doc2.paragraphs:
    assert para.style.name == 'List Bullet', f"应该是List Bullet，实际是{para.style.name}"
print("✅ 无序列表样式验证通过")

print("\n" + "=" * 60)
print("正确测试3: 验证有序列表样式（Markdown规范：自动重新编号）")
print("=" * 60)

test_file3 = temp_dir / 'test_number_list.docx'
content3 = """1. 第一项
2. 第二项
10. 第十项（Markdown规范：会重新编号为3）"""

write_docx(str(test_file3), content=content3)
doc3 = Document(str(test_file3))

assert len(doc3.paragraphs) == 3, "应该有3个段落"
for para in doc3.paragraphs:
    assert para.style.name == 'List Number', f"应该是List Number，实际是{para.style.name}"

# 验证内容（数字前缀被删除，符合Markdown规范）
assert doc3.paragraphs[0].text == "第一项", f"内容应该是'第一项'，实际是'{doc3.paragraphs[0].text}'"
assert doc3.paragraphs[1].text == "第二项", f"内容应该是'第二项'，实际是'{doc3.paragraphs[1].text}'"
assert "第十项" in doc3.paragraphs[2].text, f"内容应该包含'第十项'，实际是'{doc3.paragraphs[2].text}'"
print("✅ 有序列表样式验证通过")
print("  注意：数字前缀被删除，Word会自动重新编号为1,2,3")

print("\n" + "=" * 60)
print("正确测试4: 验证混合内容")
print("=" * 60)

test_file4 = temp_dir / 'test_mixed.docx'
content4 = """# 标题

普通段落

- 列表项1
- 列表项2"""

write_docx(str(test_file4), content=content4)
doc4 = Document(str(test_file4))

# 验证混合内容
assert doc4.paragraphs[0].style.name == 'Heading 1', "第一个应该是标题"
assert doc4.paragraphs[1].style.name == 'Normal', "第二个应该是普通段落"
assert doc4.paragraphs[2].style.name == 'List Bullet', "第三个应该是列表"
print("✅ 混合内容验证通过")

print("\n" + "=" * 60)
print("所有验证通过！")
print("=" * 60)

# 清理
shutil.rmtree(temp_dir)