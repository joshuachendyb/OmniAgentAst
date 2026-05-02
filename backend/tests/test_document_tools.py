# -*- coding: utf-8 -*-
"""
文档读写工具测试模块

【创建时间】2026-05-02 小沈

测试覆盖：
- read_pdf: PDF文件读取
- read_docx: Word文档读取
- read_xlsx: Excel文件读取

Author: 小沈 - 2026-05-02
"""

import os
import sys
import tempfile
import pytest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from app.services.tools.document.document_tools import (
    read_pdf,
    read_docx,
    read_xlsx,
)


class TestReadPdf:
    """read_pdf 工具测试 - 小沈 2026-05-02"""

    def test_no_pdfplumber(self):
        try:
            import pdfplumber
            pytest.skip("pdfplumber installed, cannot test missing case")
        except ImportError:
            result = read_pdf(file_path="D:/nonexistent.pdf")
            assert result["code"] == "ERR_NO_PDFPLUMBER"

    def test_file_not_exists(self):
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber not installed")

        result = read_pdf(file_path="D:/nonexistent/file.pdf")
        assert result["code"] == "ERR_READ_PDF"

    def test_read_basic_pdf(self):
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber not installed")

        fd, path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(path)
            c.drawString(100, 750, "Hello PDF Test")
            c.save()

            result = read_pdf(file_path=path)
            assert result["code"] == "SUCCESS"
            assert result["data"]["page_count"] >= 1
        except ImportError:
            pytest.skip("reportlab not installed for creating test PDF")
        finally:
            if os.path.exists(path):
                os.unlink(path)


class TestReadDocx:
    """read_docx 工具测试 - 小沈 2026-05-02"""

    def test_no_python_docx(self):
        try:
            import docx
            pytest.skip("python-docx installed, cannot test missing case")
        except ImportError:
            result = read_docx(file_path="D:/nonexistent.docx")
            assert result["code"] == "ERR_NO_DOCX"

    def test_file_not_exists(self):
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")

        result = read_docx(file_path="D:/nonexistent/file.docx")
        assert result["code"] == "ERR_READ_DOCX"

    def test_read_basic_docx(self):
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc = docx.Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("Second paragraph")
            doc.save(path)

            result = read_docx(file_path=path)
            assert result["code"] == "SUCCESS"
            assert result["data"]["paragraph_count"] == 2
            assert "Hello World" in result["data"]["text"]
        finally:
            if os.path.exists(path):
                os.unlink(path)


class TestReadXlsx:
    """read_xlsx 工具测试 - 小沈 2026-05-02"""

    def test_no_openpyxl(self):
        try:
            import openpyxl
            pytest.skip("openpyxl installed, cannot test missing case")
        except ImportError:
            result = read_xlsx(file_path="D:/nonexistent.xlsx")
            assert result["code"] == "ERR_NO_OPENPYXL"

    def test_file_not_exists(self):
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")

        result = read_xlsx(file_path="D:/nonexistent/file.xlsx")
        assert result["code"] == "ERR_READ_XLSX"

    def test_read_basic_xlsx(self):
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        fd, path = tempfile.mkstemp(suffix=".xlsx")
        os.close(fd)
        try:
            wb = Workbook()
            ws = wb.active
            ws.append(["name", "age", "score"])
            ws.append(["Alice", 25, 90])
            ws.append(["Bob", 30, 85])
            wb.save(path)

            result = read_xlsx(file_path=path)
            assert result["code"] == "SUCCESS"
            assert result["data"]["row_count"] == 2
            assert "name" in result["data"]["headers"]
        finally:
            if os.path.exists(path):
                os.unlink(path)

    def test_read_xlsx_with_max_rows(self):
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        fd, path = tempfile.mkstemp(suffix=".xlsx")
        os.close(fd)
        try:
            wb = Workbook()
            ws = wb.active
            ws.append(["name", "value"])
            for i in range(100):
                ws.append([f"item_{i}", i])
            wb.save(path)

            result = read_xlsx(file_path=path, max_rows=10)
            assert result["code"] == "SUCCESS"
            assert result["data"]["row_count"] == 10
        finally:
            if os.path.exists(path):
                os.unlink(path)

    def test_read_xlsx_sheet_not_found(self):
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        fd, path = tempfile.mkstemp(suffix=".xlsx")
        os.close(fd)
        try:
            wb = Workbook()
            wb.save(path)

            result = read_xlsx(file_path=path, sheet_name="NonExistent")
            assert result["code"] == "ERR_READ_XLSX"
        finally:
            if os.path.exists(path):
                os.unlink(path)
